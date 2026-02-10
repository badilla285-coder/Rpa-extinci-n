import streamlit as st
from docx import Document
from docx.shared import Pt
import PyPDF2
import io

def crear_escrito(datos, texto_condena):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("SUMILLA: SOLICITA EXTINCIÓN RPA.\n").bold = True
    p.add_run(f"RIT: {datos['rit']}\n")
    p.add_run(f"RUC: {datos['ruc']}\n")
    p.add_run(f"TRIBUNAL: {datos['juzgado']}\n")
    doc.add_paragraph("\nEN LO PRINCIPAL: SOLICITA EXTINCIÓN; OTROSÍ: ACOMPAÑA.")
    p_juez = doc.add_paragraph()
    p_juez.add_run(f"\nS.J.L. DE GARANTÍA DE {datos['juzgado'].upper()}").bold = True
    cuerpo = doc.add_paragraph()
    cuerpo.add_run(f"\n{datos['nombre_defensor']}, defensor de {datos['nombre_adolescente']}, causa RIT {datos['rit']}, digo:\n")
    cuerpo.add_run("\nQue, por este acto y según Ley 20.084, solicito declarar la extinción de responsabilidad penal de mi representado, fundado en condena de adulto privativa de libertad, incompatible con el sistema RPA.\n")
    doc.add_paragraph("\nFUNDAMENTOS TRANSCRITOS:").bold = True
    doc.add_paragraph(texto_condena)
    p_final = doc.add_paragraph()
    p_final.add_run("\nPOR TANTO, según Ley 20.084 y normas de extinción:\n")
    p_final.add_run("SOLICITO A SS. declarar la extinción y el archivo.").bold = True
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

st.set_page_config(page_title="Generador RPA")
st.title("⚖️ Generador de Extinciones")
with st.form("f1"):
    c1, c2 = st.columns(2)
    with c1:
        nom_def = st.text_input("Defensor")
        nom_ado = st.text_input("Adolescente")
        juz = st.text_input("Juzgado")
    with c2:
        ruc = st.text_input("RUC")
        rit = st.text_input("RIT")
        t_con = st.text_input("Tribunal Adulto")
    pdf = st.file_uploader("PDF Condena", type="pdf")
    btn = st.form_submit_button("Generar")
if btn and pdf:
    try:
        reader = PyPDF2.PdfReader(pdf)
        txt = ""
        for page in reader.pages:
            txt += page.extract_text() + "\n"
        d_c = {"nombre_defensor":nom_def,"nombre_adolescente":nom_ado,"juzgado":juz,"ruc":ruc,"rit":rit,"condenado_en":t_con}
        res = crear_escrito(d_c, txt)
        st.success("Hecho.")
        st.download_button("📥 Descargar Word", res, f"Extincion_{nom_ado}.docx")
    except Exception as e:
        st.error(f"Error: {e}")
