import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import io
import re
import json
from datetime import datetime
import PyPDF2
from supabase import create_client
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import time
import random
import tempfile
import os
import numpy as np # Importante para los vectores
# ... (Tus imports existentes hasta numpy)
import numpy as np

# === NUEVOS IMPORTS LANGCHAIN (ESTRUCTURA ROBUSTA) ===
try:
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.prompts import PromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.runnables import RunnablePassthrough
except ImportError:
    # Fallback para versiones antiguas o instalaciones híbridas
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain.prompts import PromptTemplate
    from langchain.schema.output_parser import StrOutputParser
    from langchain.schema.runnable import RunnablePassthrough

# Configuración de LangChain con Gemini
def get_langchain_model():
    """Retorna una instancia de ChatGoogleGenerativeAI configurada."""
    try:
        # Verificamos si la API KEY existe en secrets para evitar errores mudos
        if "GOOGLE_API_KEY" not in st.secrets:
            st.error("Falta la configuración de GOOGLE_API_KEY en los Secrets de Streamlit.")
            return None
            
        api_key = st.secrets["GOOGLE_API_KEY"]
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash", 
            temperature=0.3, # Precisión legal mantenida
            google_api_key=api_key,
            # Esta opción ayuda a la compatibilidad con modelos antiguos si fuera necesario
            convert_system_message_to_human=True 
        )
        return llm
    except Exception as e:
        st.error(f"Error iniciando LangChain: {e}")
        return None
        # === CONFIGURACIÓN DE PROMPT TEMPLATE (ESTÁNDAR JURÍDICO) ===

def get_legal_prompt_template():
    """
    Define la estructura de pensamiento del modelo.
    Diseñado para análisis de documentos y asistencia legal.
    """
    template = """
    Eres un asistente de Inteligencia Artificial especializado en el ámbito jurídico chileno, 
    actuando como un apoyo experto para un profesional del derecho.

    CONTEXTO DEL CASO:
    {context}

    INSTRUCCIÓN ESPECÍFICA:
    {question}

    DIRECTRICES DE RESPUESTA:
    1. Utiliza terminología técnica precisa (ej. "procedimiento ordinario", "excepciones dilatorias", "acción de protección").
    2. Si el contexto no contiene información suficiente, admítelo con honestidad.
    3. Organiza la respuesta de forma jerárquica y clara.
    4. Mantén un tono formal, objetivo y analítico.

    RESPUESTA TÉCNICA:
    """
    
    return PromptTemplate(
        template=template,
        input_variables=["context", "question"]
    )

# === INTEGRACIÓN EN LA CADENA DE TRABAJO (CHAIN) ===

def process_legal_query(user_question, context_data):
    """
    Ejecuta la cadena de LangChain (Chain) para procesar la consulta.
    """
    llm = get_langchain_model()
    if llm is None:
        return "Error: No se pudo inicializar el modelo."

    prompt = get_legal_prompt_template()
    
    # Construcción de la cadena usando el lenguaje de expresión de LangChain (LCEL)
    chain = (
        {"context": lambda x: context_data, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    try:
        response = chain.invoke(user_question)
        return response
    except Exception as e:
        return f"Error durante el procesamiento: {str(e)}"
        # --- FUNCIONES DE EXTRACCIÓN DE TEXTO ---

def extraer_texto_pdf(file):
    """Extrae texto de un archivo PDF."""
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extraer_texto_docx(file):
    """Extrae texto de un archivo Word."""
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extraer_texto_generico(uploaded_file):
    """Detecta el tipo de archivo y extrae el texto de forma automática."""
    if uploaded_file is None: return ""
    try:
        if uploaded_file.name.lower().endswith('.pdf'):
            return extraer_texto_pdf(uploaded_file)
        elif uploaded_file.name.lower().endswith(('.docx', '.doc')):
            return extraer_texto_docx(uploaded_file)
        return ""
    except Exception as e:
        st.error(f"Error extrayendo texto de {uploaded_file.name}: {e}")
        return ""
# ... (Continúa con tu código de configuración y CSS)

# =============================================================================
# 1. CONFIGURACIÓN Y ESTILOS (INTERFAZ ELEGANTE & LEGIBLE)
# =============================================================================
st.set_page_config(
    page_title="Sistema Jurídico Avanzado IABL",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Profesional: Diseño "LegalTech Premium"
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Manrope:wght@300;400;600;800&display=swap');
    
    /* === TIPOGRAFÍA Y FONDO === */
    .stApp {
        background-color: #F8F9FA; /* Blanco humo muy suave */
        font-family: 'Manrope', sans-serif;
    }
    
    /* === SIDEBAR (BARRA LATERAL) === */
    section[data-testid="stSidebar"] {
        background-color: #161B2F; /* Navy Profundo */
        box-shadow: 2px 0 10px rgba(0,0,0,0.1);
    }
    section[data-testid="stSidebar"] h1, 
    section[data-testid="stSidebar"] h2, 
    section[data-testid="stSidebar"] h3, 
    section[data-testid="stSidebar"] label, 
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] div {
        color: #ECEFF1 !important; /* Blanco hueso para texto sidebar */
    }
    /* Separador en sidebar */
    section[data-testid="stSidebar"] hr {
        border-color: #5B687C !important;
    }

    /* === ENCABEZADOS PRINCIPALES === */
    h1 { 
        color: #161B2F !important; 
        font-weight: 800; 
        letter-spacing: -0.02em;
        border-bottom: 3px solid #D4CDCB;
        padding-bottom: 15px; 
        margin-bottom: 25px;
        font-size: 2.2rem;
    }
    h2 { color: #2C3550 !important; font-weight: 700; margin-top: 20px; }
    h3 { color: #5B687C !important; font-weight: 600; }
    
    /* === BOTONES (Elegantes y con feedback) === */
    .stButton>button {
        background: linear-gradient(135deg, #161B2F 0%, #2C3550 100%) !important;
        color: white !important;
        border-radius: 8px;
        border: none;
        padding: 0.6rem 1.2rem;
        font-weight: 600;
        letter-spacing: 0.5px;
        transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
        box-shadow: 0 4px 6px rgba(22, 27, 47, 0.1);
        width: 100%;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 7px 14px rgba(22, 27, 47, 0.2);
        background: linear-gradient(135deg, #2C3550 0%, #161B2F 100%) !important;
    }
    /* Botones secundarios (tipo "Limpiar") */
    button[kind="secondary"] {
        background: transparent !important;
        border: 2px solid #5B687C !important;
        color: #5B687C !important;
    }
    button[kind="secondary"]:hover {
        background: #ECEFF1 !important;
    }

    /* === INPUTS Y TEXTAREAS === */
    div[data-baseweb="input"] > div, 
    div[data-baseweb="textarea"] > div,
    div[data-baseweb="select"] > div {
        background-color: #FFFFFF;
        border-radius: 8px;
        border: 1px solid #D4CDCB;
        color: #161B2F;
    }
    div[data-baseweb="input"]:focus-within > div {
        border-color: #161B2F !important;
        box-shadow: 0 0 0 2px rgba(22, 27, 47, 0.1);
    }

    /* === TABS (PESTAÑAS) === */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        border-bottom: 1px solid #D4CDCB;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        border-radius: 8px 8px 0 0;
        background-color: transparent;
        border: none;
        color: #5B687C;
        font-weight: 600;
    }
    .stTabs [data-baseweb="tab"]:hover {
        color: #161B2F;
        background-color: #ECEFF1;
    }
    .stTabs [aria-selected="true"] {
        background-color: white !important;
        color: #161B2F !important;
        border-bottom: 3px solid #161B2F !important;
    }

    /* === CAJAS Y TARJETAS === */
    .status-card {
        padding: 20px;
        border-radius: 12px;
        background: #ffffff;
        border-left: 6px solid #161B2F;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        color: #212121;
        margin-bottom: 20px;
    }
    
    /* Minuta / Documento */
    .minuta-box {
        background-color: #fff;
        padding: 40px;
        border-radius: 4px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.08);
        font-family: 'Courier New', Courier, monospace; 
        color: #333;
        position: relative;
    }
    .minuta-box::before {
        content: '';
        position: absolute;
        top: 0; left: 0; bottom: 0;
        width: 6px;
        background: #FBC02D;
    }
    
    /* Resumen IA y Badges */
    .resumen-dinamico {
        background: linear-gradient(to right, #ffffff, #fcfcfc);
        border: 1px solid #e0e0e0;
        border-left: 5px solid #2C3550;
        padding: 25px;
        border-radius: 12px;
        margin-bottom: 20px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.03);
    }
    .badge-tipo {
        background-color: #161B2F;
        color: #fff;
        padding: 4px 10px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 700;
        letter-spacing: 0.5px;
        text-transform: uppercase;
    }
    .badge-rol {
        background-color: #ECEFF1;
        color: #455A64;
        padding: 4px 10px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 700;
        border: 1px solid #CFD8DC;
        margin-left: 5px;
    }
    </style>
""", unsafe_allow_html=True)

# =============================================================================
# 2. CONFIGURACIÓN SERVICIOS (SEGURIDAD REFORZADA)
# =============================================================================

# === CONFIGURACIÓN SEGURA (SECRETS) ===
try:
    if "GOOGLE_API_KEY" in st.secrets:
        genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    else:
        st.error("⚠️ FALTA CONFIGURAR LA API KEY EN SECRETS (GOOGLE_API_KEY).")
except Exception as e:
    st.error(f"⚠️ Error configurando API Key: {e}")

# === NUEVA FUNCIÓN MAESTRA DE MODELOS DINÁMICOS ===
def get_generative_model_dinamico():
    """Busca automáticamente un modelo generativo disponible (Flash > Pro > Cualquiera)."""
    try:
        modelos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        # Prioridad: 1.5 Flash -> 1.5 Pro -> Cualquiera
        mejor = next((m for m in modelos if 'gemini-1.5-flash' in m), None)
        if not mejor:
            mejor = next((m for m in modelos if 'gemini-1.5-pro' in m), modelos[0])
        return genai.GenerativeModel(mejor)
    except Exception as e:
        # Fallback de emergencia por si la lista falla
        return genai.GenerativeModel('models/gemini-1.5-flash-latest')

# Instancia global inicial
model_ia = get_generative_model_dinamico()

# === LÓGICA DE DETECCIÓN AUTOMÁTICA DE MODELO DE EMBEDDING ===
MODELO_EMBEDDING_ACTUAL = None

def get_embedding_model():
    """Busca automáticamente un modelo de embedding disponible en la cuenta."""
    global MODELO_EMBEDDING_ACTUAL
    if MODELO_EMBEDDING_ACTUAL:
        return MODELO_EMBEDDING_ACTUAL

    try:
        modelos = list(genai.list_models())
        for m in modelos:
            if 'embedContent' in m.supported_generation_methods:
                if 'text-embedding-004' in m.name:
                    MODELO_EMBEDDING_ACTUAL = m.name
                    return m.name
        for m in modelos:
            if 'embedContent' in m.supported_generation_methods:
                MODELO_EMBEDDING_ACTUAL = m.name
                return m.name
        return 'models/text-embedding-004'
    except Exception as e:
        return 'models/text-embedding-004'

# === FUNCIÓN PARA METADATA PROFUNDA (ACTUALIZADA LISTA CERRADA) ===
def analizar_metadata_profunda(texto_completo):
    """Usa IA para extraer metadata precisa del texto completo del documento."""
    try:
        prompt = f"""
        Eres un Actuario Judicial experto. Lee este documento legal COMPLETO. 
        Extrae con precisión quirúrgica un JSON válido con los siguientes campos.
        
        IMPORTANTE: El campo 'tipo' debe ser ESTRICTAMENTE uno de estos valores:
        ["Sentencia Condenatoria", "Sentencia Absolutoria", "Recurso de Nulidad", "Recurso de Amparo", "Recurso de Apelación", "Doctrina/Artículo", "Ley/Normativa"].
        
        JSON REQUERIDO:
        {{
            "tribunal": "Nombre exacto del tribunal (ej: Corte de Apelaciones de Santiago)",
            "rol": "RIT o Rol de la causa (ej: 450-2023)",
            "fecha_sentencia": "Fecha del documento o sentencia (YYYY-MM-DD) o 'S/F'",
            "resultado": "Resumen muy breve (ej: Acoge Recurso, Rechaza Nulidad)",
            "tema": "Palabras clave del tema jurídico (ej: Indicios Art 85, Prisión Preventiva)",
            "tipo": "Uno de la lista cerrada anterior"
        }}
        
        TEXTO DEL DOCUMENTO (Primeros 20000 caracteres):
        {texto_completo[:20000]}
        """
        
        model = get_generative_model_dinamico()
        resp = model.generate_content(prompt)
        clean_json = resp.text.replace('```json', '').replace('```', '').strip()
        return json.loads(clean_json)
    except Exception as e:
        return {
            "tribunal": "Desconocido/Error IA",
            "rol": "S/N",
            "fecha_sentencia": datetime.now().strftime("%Y-%m-%d"),
            "resultado": "Pendiente",
            "tema": "General",
            "tipo": "Documento Legal"
        }

# === INICIALIZACIÓN SEGURA DE SUPABASE (SECRETOS) ===
@st.cache_resource
def init_supabase():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except Exception as e:
        st.error("⚠️ Error de configuración: Faltan las claves de Supabase en Secrets (.streamlit/secrets.toml).")
        return None

supabase = init_supabase()

# =============================================================================
# 3. DATOS MAESTROS Y LÓGICA PENAL MATEMÁTICA
# =============================================================================
TRIBUNALES = [
    "1° Juzgado de Garantía de Santiago", "2° Juzgado de Garantía de Santiago",
    "3° Juzgado de Garantía de Santiago", "4° Juzgado de Garantía de Santiago",
    "5° Juzgado de Garantía de Santiago", "6° Juzgado de Garantía de Santiago",
    "7° Juzgado de Garantía de Santiago", "8° Juzgado de Garantía de Santiago",
    "9° Juzgado de Garantía de Santiago", "Juzgado de Garantía de San Bernardo", 
    "Juzgado de Garantía de Puente Alto", "Juzgado de Garantía de Talagante", 
    "Juzgado de Garantía de Melipilla", "Juzgado de Garantía de Colina",
    "3° Tribunal de Juicio Oral en lo Penal de Santiago",
    "Iltma. Corte de Apelaciones de San Miguel",
    "Iltma. Corte de Apelaciones de Santiago"
]

TIPOS_RECURSOS = [
    "Extinción Art. 25 ter",
    "Prescripción de la Pena",
    "Amparo Constitucional",
    "Apelación por Quebrantamiento"
]

# Escala de Penas (Grados) para cálculo matemático
ESCALA_PENAS = [
    {"nombre": "Prisión en su grado mínimo", "min": 1, "max": 20},
    {"nombre": "Prisión en su grado medio", "min": 21, "max": 40},
    {"nombre": "Prisión en su grado máximo", "min": 41, "max": 60},
    {"nombre": "Presidio menor en su grado mínimo", "min": 61, "max": 540},
    {"nombre": "Presidio menor en su grado medio", "min": 541, "max": 1095}, # 3 años
    {"nombre": "Presidio menor en su grado máximo", "min": 1096, "max": 1825}, # 5 años
    {"nombre": "Presidio mayor en su grado mínimo", "min": 1826, "max": 3650}, # 10 años
    {"nombre": "Presidio mayor en su grado medio", "min": 3651, "max": 5475}, # 15 años
    {"nombre": "Presidio mayor en su grado máximo", "min": 5476, "max": 7300}, # 20 años
    {"nombre": "Presidio perpetuo", "min": 7301, "max": 14600} # Simbólico
]

# Base de datos de delitos con índice de grado base en ESCALA_PENAS
DELITOS_INFO = {
    "Robo con Intimidación": {"idx_min": 6, "idx_max": 8},
    "Robo con Violencia": {"idx_min": 6, "idx_max": 8},
    "Robo en Lugar Habitado": {"idx_min": 6, "idx_max": 6},
    "Microtráfico (Art. 4)": {"idx_min": 4, "idx_max": 5},
    "Tráfico Ilícito (Art. 3)": {"idx_min": 6, "idx_max": 7},
    "Homicidio Simple": {"idx_min": 7, "idx_max": 8},
    "Receptación": {"idx_min": 3, "idx_max": 5},
    "Porte Ilegal de Arma": {"idx_min": 5, "idx_max": 6},
    "Lesiones Graves": {"idx_min": 4, "idx_max": 4},
    "Amenazas Simples": {"idx_min": 3, "idx_max": 3},
    "Maltrato de Obra a Carabineros": {"idx_min": 4, "idx_max": 5}
}

# =============================================================================
# 4. LÓGICA DE IA & PROCESAMIENTO
# =============================================================================
def analizar_pdf(uploaded_file, tipo):
    if not model_ia: return None
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        text = "".join([page.extract_text() for page in reader.pages[:3]])
        prompt = f"""
        Analiza este documento legal chileno ({tipo}). Extrae en JSON:
        {{
            "rit": "RIT completo", "ruc": "RUC completo",
            "tribunal": "Nombre tribunal", "imputado": "Nombre completo",
            "fecha_sentencia": "YYYY-MM-DD", "pena": "Texto de la pena",
            "sancion": "Texto de sanción RPA"
        }}
        Texto: {text[:4000]}
        """
        resp = model_ia.generate_content(prompt)
        clean_json = resp.text.replace('```json', '').replace('```', '').strip()
        return json.loads(clean_json)
    except Exception as e:
        st.error(f"Error IA: {e}")
        return None

# =============================================================================
# 5. MOTOR DE GENERACIÓN WORD
# =============================================================================
class GeneradorWord:
    def __init__(self, defensor, imputado):
        self.doc = Document()
        self.defensor = defensor.upper() if defensor else "DEFENSOR PÚBLICO"
        self.imputado = imputado.upper() if imputado else "IMPUTADO"
        
        section = self.doc.sections[0]
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(12)
        
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_parrafo(self, texto, negrita=False, align="JUSTIFY", sangria=True):
        p = self.doc.add_paragraph()
        
        if align == "CENTER": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "LEFT": p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if sangria and align == "JUSTIFY":
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        texto_final = texto.replace("{DEFENSOR}", self.defensor).replace("{IMPUTADO}", self.imputado)
        
        if negrita:
            run = p.add_run(texto_final)
            run.font.name = 'Cambria'
            run.font.size = Pt(12)
            run.bold = True
        else:
            keywords = [
                r"RIT:?\s?[\w\d-]+", r"RUC:?\s?[\w\d-]+", 
                "POR TANTO", "OTROSÍ", "EN LO PRINCIPAL", 
                "SOLICITA", "INTERPONE", "ACCIÓN CONSTITUCIONAL",
                "HECHOS:", "DERECHO:", "AGRAVIO:", "PETICIONES CONCRETAS:", 
                "FUNDAMENTOS DE DERECHO:", "ANTECEDENTES DE HECHO:",
                "RESOLUCIÓN IMPUGNADA:", "ARGUMENTOS DE LA DEFENSA:", "ANTECEDENTES SOCIALES:", "SANCIÓN:", "SANCIÓN QUEBRANTADA:"
            ]
            
            patron_regex = "|".join(keywords) + f"|{re.escape(self.defensor)}|{re.escape(self.imputado)}"
            matches = list(re.finditer(patron_regex, texto_final, flags=re.IGNORECASE))
            
            last_pos = 0
            for match in matches:
                start, end = match.span()
                if start > last_pos:
                    run = p.add_run(texto_final[last_pos:start])
                    run.font.name = 'Cambria'
                    run.font.size = Pt(12)
                
                run_bold = p.add_run(texto_final[start:end])
                run_bold.font.name = 'Cambria'
                run_bold.font.size = Pt(12)
                run_bold.bold = True
                last_pos = end
            
            if last_pos < len(texto_final):
                run = p.add_run(texto_final[last_pos:])
                run.font.name = 'Cambria'
                run.font.size = Pt(12)

    def generar(self, tipo, datos):
        # 1. ENCABEZADO
        sumas = {
            "Extinción Art. 25 ter": "EN LO PRINCIPAL: SOLICITA EXTINCIÓN; OTROSÍ: ACOMPAÑA DOCUMENTO.",
            "Prescripción de la Pena": "EN LO PRINCIPAL: Solicita Audiencia de Prescripción; OTROSÍ: Oficia a extranjería y se remita extracto de filiación y antecedentes.",
            "Amparo Constitucional": "EN LO PRINCIPAL: ACCIÓN CONSTITUCIONAL DE AMPARO; OTROSÍ: ORDEN DE NO INNOVAR.",
            "Apelación por Quebrantamiento": "EN LO PRINCIPAL: INTERPONE RECURSO DE APELACIÓN; OTROSÍ: FORMA DE NOTIFICACIÓN."
        }
        self.add_parrafo(sumas.get(tipo, "SOLICITUD"), negrita=True, align="LEFT", sangria=False)
        self.doc.add_paragraph() 

        # 2. TRIBUNAL
        destinatario = "ILTMA. CORTE DE APELACIONES DE SANTIAGO" if tipo in ["Amparo Constitucional", "Apelación por Quebrantamiento"] else datos.get('tribunal_ej', 'TRIBUNAL').upper()
        self.add_parrafo(destinatario, negrita=True, align="CENTER", sangria=False)
        self.doc.add_paragraph()

        # 3. COMPARECENCIA (Multicausa)
        causas_str = ""
        lista_ind = datos.get('lista_individualizacion', [])
        if lista_ind:
            causas_txts = [f"RUC {c['ruc']}, RIT {c['rit']}" for c in lista_ind if c['ruc']]
            if causas_txts:
                causas_str = ", en las causas " + "; ".join(causas_txts) + ","
        
        elif tipo == "Prescripción de la Pena":
            lista_causas = datos.get('prescripcion_list', [])
            causas_txts = [f"RUC {c['ruc']}, RIT {c['rit']}" for c in lista_causas if c['ruc']]
            if causas_txts:
                causas_str = ", en las causas " + "; ".join(causas_txts) + ","
        elif tipo == "Apelación por Quebrantamiento":
            rit_ap = datos.get('rit_ap', '')
            ruc_ap = datos.get('ruc_ap', '')
            if rit_ap:
                causas_str = f", en causa RIT {rit_ap}, RUC {ruc_ap},"
        else:
            lista_ej = datos.get('ejecucion', [])
            causas_txts = [f"RUC {c.get('ruc','')}, RIT {c.get('rit','')}" for c in lista_ej if c.get('rit')]
            if causas_txts and not causas_str:
                causas_str = ", en causas " + "; ".join(causas_txts) + ","

        intro = f"{{DEFENSOR}}, Abogada, Defensora Penal Pública, en representación de {{IMPUTADO}}{causas_str} a S.S. respetuosamente digo:"
        self.add_parrafo(intro)

        # 4. CUERPO DEL ESCRITO
        if tipo == "Prescripción de la Pena":
            self.add_parrafo("Que, por medio de la presente, vengo en solicitar a S.S. se sirva fijar día y hora para celebrar audiencia con el objeto de debatir sobre la prescripción de la pena respecto de mi representado, de conformidad a lo dispuesto en el artículo 5 de la Ley N° 20.084 y las normas pertinentes del Código Penal.")
            self.add_parrafo("Fundamento esta solicitud en que existen sentencias condenatorias en las causas señaladas, cuyo cumplimiento a la fecha se encuentra prescrito por el transcurso del tiempo, conforme a los siguientes antecedentes:")
            lista_p = datos.get('prescripcion_list', [])
            if not lista_p:
                self.add_parrafo("(Debe ingresar las causas en el formulario lateral)")
            for c in lista_p:
                parrafo_causa = (
                    f"En la causa RUC {c['ruc']} (RIT {c['rit']} de este Tribunal): Mi representado fue condenado por sentencia de fecha {c['fecha_sentencia']}, "
                    f"dictada por el {c['tribunal']} a la pena de {c['pena']} por el delito de {c['delito']}. "
                    f"Dicha sentencia se encuentra ejecutoriada (o con cumplimiento suspendido) desde el {c['fecha_suspension']}."
                )
                self.add_parrafo(parrafo_causa)
            self.add_parrafo("Teniendo presente el tiempo transcurrido desde las fechas de las sentencias y, específicamente, desde la suspensión del cumplimiento, hasta la fecha actual (transcurriendo en exceso el plazo legal exigido para la prescripción de las sanciones en el marco de la Responsabilidad Penal Adolescente), solicito se fije audiencia con el objeto de debatir y declarar la prescripción de la pena y el consecuente sobreseimiento definitivo.")
            self.add_parrafo("POR TANTO, en mérito de lo expuesto y normativa legal citada,", sangria=False)
            self.add_parrafo("SOLICITO A S. S. acceder a lo solicitado, fijando día y hora para celebrar audiencia a fin de que se abra debate y se declare la prescripción de las penas en las presentes causas.", sangria=False)
            self.add_parrafo("OTROSÍ: Que, de conformidad a la petición principal planteada y para contar con todos los antecedentes necesarios para la adecuada resolución del tribunal, vengo en solicitar a S. S. se oficie a Extranjería con el fin de que informen los movimientos migratorios de mi representado {IMPUTADO}, desde la fecha de la primera sentencia hasta la fecha actual. Asimismo, solicito que se requiera y se incorpore a la carpeta digital el Extracto de Filiación y Antecedentes actualizado.", negrita=False)
            self.add_parrafo("POR TANTO,", sangria=False)
            self.add_parrafo("SOLICITO A S. S. acceder a lo solicitado, oficiando a Extranjería y ordenando la remisión del extracto de filiación y antecedentes actualizado.", sangria=False)

        elif tipo == "Extinción Art. 25 ter":
            self.add_parrafo("Que, vengo en solicitar que declare la extinción de las sanciones de la Ley de Responsabilidad Penal Adolescente, o en subsidio se fije día y hora para celebrar audiencia para debatir sobre la extinción de la pena respecto de mi representado, en virtud del artículo 25 ter y 25 quinquies de la Ley 20.084.")
            self.add_parrafo("Mi representado fue condenado en la siguiente causa de la Ley RPA:")
            rpas = datos.get('rpa', [])
            for idx, rpa in enumerate(rpas, 1):
                txt = f"{idx}. RIT: {rpa.get('rit','')}, RUC: {rpa.get('ruc','')}: Condenado por el {rpa.get('tribunal','JUZGADO DE GARANTÍA')} a la pena de {rpa.get('sancion','')}, debiendo cumplirse con todas las prescripciones establecidas en la ley 20.084."
                self.add_parrafo(txt)
            self.add_parrafo("El fundamento para solicitar la discusión radica en una condena de mayor gravedad como adulto:")
            ads = datos.get('adulto', [])
            for idx, ad in enumerate(ads, 1):
                txt = f"{idx}. RIT: {ad.get('rit','')}, RUC: {ad.get('ruc','')}: Condenado por el {ad.get('tribunal','')} con fecha {ad.get('fecha','')}, a la pena de {ad.get('pena','')}, como autor de delito."
                self.add_parrafo(txt)
            self.add_parrafo("Se hace presente que el artículo 25 ter en su inciso tercero establece que se considerará más grave el delito o conjunto de ellos que tuviere asignada en la ley una mayor pena de conformidad con las reglas generales.")
            self.add_parrafo("POR TANTO,", sangria=False)
            self.add_parrafo("En mérito de lo expuesto, SOLICITO A S.S. acceder a lo solicitado extinguiendo de pleno derecho la sanción antes referida.", sangria=False)
            self.add_parrafo("OTROSÍ: Acompaña sentencia de adulto.", negrita=True, sangria=False)
            self.add_parrafo("POR TANTO, SOLICITO A S.S. se tenga por acompañada.", sangria=False)

        elif tipo == "Amparo Constitucional":
            self.add_parrafo("Que, en virtud de lo dispuesto en el artículo 21 de la Constitución Política de la República, vengo en deducir acción constitucional de amparo a favor de mi representado, por la perturbación grave e ilegítima a su libertad personal y seguridad individual.")
            self.add_parrafo("ANTECEDENTES DE HECHO:", negrita=True)
            if datos.get('argumento_extra'):
                self.add_parrafo(datos['argumento_extra'])
            else:
                self.add_parrafo("La resolución recurrida ordenó el ingreso inmediato del joven, quebrantando una sanción de adolescente, la cual no se encontraba ejecutoriada y estando pendiente recurso de apelación, siendo la resolución ilegal y arbitraria.")
            self.add_parrafo("FUNDAMENTOS DE DERECHO:", negrita=True)
            self.add_parrafo("1. Normativa Internacional y Constitucional: El derecho a la libertad personal se encuentra garantizado en el artículo 7 de la Convención Americana de Derechos Humanos y el artículo 19 Nº 7 de la Constitución Política de la República. El artículo 21 de la Carta Fundamental establece el recurso de amparo como la vía idónea para restablecer el imperio del derecho.")
            self.add_parrafo("2. Vulneración del artículo 79 del Código Penal: Dicha norma establece que 'no podrá ejecutarse pena alguna sino en virtud de sentencia ejecutoriada'. En el presente caso, la resolución impugnada ordena un ingreso o mantiene una privación de libertad sin que exista una sentencia firme que lo habilite, vulnerando el principio de legalidad.")
            self.add_parrafo("3. Interés Superior del Adolescente y Convención de Derechos del Niño: El artículo 37 letra b) de la Convención prescribe que la detención o prisión de un niño se utilizará tan sólo como medida de último recurso y durante el período más breve que proceda.")
            self.add_parrafo("POR TANTO,", sangria=False)
            self.add_parrafo("SOLICITO A V.S. ILTMA. admitir a tramitación la presente acción, pedir informe urgente al recurrido y, en definitiva, acoger el amparo, dejando sin efecto la resolución impugnada y ordenando la libertad inmediata de mi representado.", sangria=False)
            self.add_parrafo("OTROSÍ: ORDEN DE NO INNOVAR.", negrita=True, sangria=False)
            self.add_parrafo("Solicito se decrete orden de no innovar para suspender los efectos de la resolución recurrida mientras se tramita la presente acción, a fin de evitar que se consolide la afectación a la libertad personal.", sangria=False)

        elif tipo == "Apelación por Quebrantamiento":
            self.add_parrafo("Que encontrándome dentro del plazo legal, vengo en interponer recurso de apelación en contra de la resolución que ordenó el quebrantamiento definitivo de la sanción de mi representado, solicitando se revoque y se mantenga la sanción original en el medio libre o se decrete un quebrantamiento parcial.")
            self.add_parrafo("I. HECHOS:", negrita=True)
            self.add_parrafo(datos.get('hechos_quebrantamiento', 'No especificados'))
            
            self.add_parrafo("RESOLUCIÓN IMPUGNADA:", negrita=True)
            self.add_parrafo(datos.get('resolucion_tribunal', 'No especificada'))
            self.add_parrafo("ARGUMENTOS DE LA DEFENSA:", negrita=True)
            self.add_parrafo(datos.get('argumentos_defensa', 'No especificados'))
            
            if datos.get('antecedentes_sociales'):
                self.add_parrafo("ANTECEDENTES SOCIALES:", negrita=True)
                self.add_parrafo(datos.get('antecedentes_sociales'))
            
            self.add_parrafo("SANCIÓN ORIGINAL:", negrita=True)
            self.add_parrafo(datos.get('sancion_orig', ''))
            self.add_parrafo("SANCIÓN QUEBRANTADA:", negrita=True)
            self.add_parrafo(datos.get('sancion_quebrantada', ''))

            self.add_parrafo("II. EL DERECHO Y AGRAVIO:", negrita=True)
            self.add_parrafo("La resolución causa agravio pues desestima que la privación de libertad es una medida de último recurso (ultima ratio) según el artículo 40 n°2 de la Convención de Derechos del Niño.")
            self.add_parrafo("Principio de Progresividad: El artículo 52 de la Ley 20.084 establece una gradualidad en las sanciones por incumplimiento. Saltar directamente al quebrantamiento definitivo vulnera este principio, interrumpiendo procesos de reinserción escolar o laboral.")
            self.add_parrafo("Reinserción Social: El fin de la pena adolescente es la prevención especial positiva. El encierro total frustra este objetivo.")
            self.add_parrafo("POR TANTO,", sangria=False)
            self.add_parrafo("SOLICITO A US. tener por interpuesto recurso de apelación, concederlo y elevar los antecedentes a la Iltma. Corte de Apelaciones para que revoque la resolución impugnada.", sangria=False)

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

# =============================================================================
# 6. LÓGICA DE SESIÓN Y USUARIOS
# =============================================================================
if "db_users" not in st.session_state:
    st.session_state.db_users = [
        {"email": "admin@iabl.cl", "pass": "admin123", "rol": "Admin", "nombre": "IGNACIO BADILLA LARA"},
        {"email": "usuario@defensoria.cl", "pass": "defensor", "rol": "User", "nombre": "DEFENSOR PÚBLICO"}
    ]

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_role" not in st.session_state:
    st.session_state.user_role = "user"
if "defensor_nombre" not in st.session_state:
    st.session_state.defensor_nombre = ""

# =============================================================================
# 7. PANTALLA DE LOGIN (REDISEÑO HERO VERTICAL)
# =============================================================================
def login_screen():
    # CSS Personalizado para Landing Page (Con la nueva paleta)
    st.markdown("""
        <style>
        /* Estilos Generales Landing */
        .stApp {
            background-color: #F4F7F6;
        }
        h1 {
            color: #161B2F !important;
            text-align: center;
            font-weight: 800;
            font-size: 3rem;
            margin-bottom: 0.5rem;
        }
        .hero-subtitle {
            color: #5B687C;
            text-align: center;
            font-size: 1.3rem;
            font-style: italic;
            margin-bottom: 2rem;
            font-weight: 400;
        }
        /* Tarjeta de Login */
        [data-testid="stForm"] {
            background-color: white;
            padding: 3rem;
            border-radius: 20px;
            box-shadow: 0 4px 20px rgba(22, 27, 47, 0.08); /* Sombra Navy sutil */
            border: 1px solid #D4CDCB; /* Beige Suave */
        }
        /* Botones */
        .stButton>button {
            background-color: #161B2F !important;
            color: white !important;
            border-radius: 10px;
            border: none;
            font-weight: 600;
            padding: 0.8rem;
            transition: all 0.3s;
        }
        .stButton>button:hover {
            opacity: 0.9;
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        /* Features */
        .feature-box {
            text-align: center;
            padding: 1.5rem;
            background: white;
            border-radius: 15px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.05);
            border: 1px solid #D4CDCB; /* Beige Suave */
            transition: transform 0.3s;
        }
        .feature-box:hover {
            transform: translateY(-5px);
            box-shadow: 0 10px 20px rgba(22, 27, 47, 0.08);
        }
        .feature-icon {
            font-size: 2.5rem;
            margin-bottom: 1rem;
            display: block;
            color: #5B687C;
        }
        .feature-text {
            color: #161B2F;
            font-weight: 700;
        }
        </style>
    """, unsafe_allow_html=True)

    # Espacio aire
    st.write("")
    st.write("")

    # HERO SECTION
    st.markdown("<h1>SISTEMA JURÍDICO IABL</h1>", unsafe_allow_html=True)
    st.markdown("<p class='hero-subtitle'>Automatización inteligente para defensores: tu tiempo, salud y excelencia profesional.</p>", unsafe_allow_html=True)

    # LOGIN SECTION (Centrado)
    c1, c2, c3 = st.columns([1, 1.5, 1])
    
    with c2:
        # Título del formulario específico solicitado (CON TILDE)
        st.markdown("<h3 style='text-align: center; color: #161B2F; margin-bottom: 20px;'>ACCESO A SISTEMA JURÍDICO AVANZADO</h3>", unsafe_allow_html=True)
        
        tab_login, tab_registro = st.tabs(["🔐 Iniciar Sesión", "📝 Crear Cuenta"])
        
        with tab_login:
            with st.form("login_form"):
                email = st.text_input("Correo Electrónico", placeholder="usuario@defensoria.cl")
                password = st.text_input("Contraseña", type="password", placeholder="••••••••")
                st.write("") # Espaciador
                submitted = st.form_submit_button("INGRESAR AL SISTEMA", use_container_width=True)
                
                if submitted:
                    try:
                        session = supabase.auth.sign_in_with_password({"email": email, "password": password})
                        user = session.user
                        data = supabase.table("profiles").select("*").eq("id", user.id).execute()
                        if data.data:
                            perfil = data.data[0]
                            st.session_state.logged_in = True
                            st.session_state.user_role = perfil['rol']
                            st.session_state.defensor_nombre = perfil['nombre']
                            st.session_state.user_email = email
                            st.success("¡Bienvenido!")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Error: Usuario autenticado pero sin perfil.")
                    except Exception as e:
                        st.error(f"Credenciales incorrectas o error de conexión: {e}")

        with tab_registro:
            with st.form("register_form"):
                new_email = st.text_input("Tu Correo")
                new_pass = st.text_input("Crear Contraseña", type="password")
                new_name = st.text_input("Nombre Completo")
                st.write("")
                reg_submit = st.form_submit_button("REGISTRARSE", use_container_width=True)
                
                if reg_submit:
                    try:
                        response = supabase.auth.sign_up({
                            "email": new_email, 
                            "password": new_pass,
                            "options": {"data": {"nombre": new_name}}
                        })
                        st.success("✅ Cuenta creada. Revisa tu correo o intenta iniciar sesión.")
                    except Exception as e:
                        st.error(f"Error al registrar: {e}")

    # FEATURES SECTION
    st.markdown("---")
    st.write("")
    
    f1, f2, f3, f4 = st.columns(4)
    
    def feature_card(icon, title, desc):
        return f"""
        <div class='feature-box'>
            <span class='feature-icon'>{icon}</span>
            <div class='feature-text'>{title}</div>
            <div style='color: #5B687C; font-size: 0.9rem;'>{desc}</div>
        </div>
        """

    with f1:
        st.markdown(feature_card("📝", "Redacción", "Escritos automáticos"), unsafe_allow_html=True)
    with f2:
        st.markdown(feature_card("👁️", "Visión IA", "OCR Multimodal"), unsafe_allow_html=True)
    with f3:
        st.markdown(feature_card("📚", "Biblioteca", "Jurisprudencia RAG"), unsafe_allow_html=True)
    with f4:
        st.markdown(feature_card("🎙️", "Audio", "Transcripción Forense"), unsafe_allow_html=True)

# =============================================================================
# 8. CÁLCULO PENAL AVANZADO (LÓGICA JURÍDICA MATEMÁTICA)
# =============================================================================
def init_session_data():
    defaults = {
        "imputado": "", 
        "tribunal_sel": TRIBUNALES[9] if TRIBUNALES else "",
        "ejecucion": [{"rit": "", "ruc": ""}],
        "rpa": [{"rit": "", "ruc": "", "tribunal": "", "sancion": ""}],
        "adulto": [],
        "prescripcion_list": [],
        "lista_individualizacion": [],
        "all_text": "",
        "logs": []  # <--- ESTA LÍNEA ES VITAL PARA EL ERROR DE ATTRIBUTEERROR
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def calcular_pena_exacta(delito_info, atenuantes, agravantes, es_rpa):
    idx_min = delito_info["idx_min"]
    idx_max = delito_info["idx_max"]
    
    n_at = len(atenuantes)
    n_ag = len(agravantes)
    
    if n_at > 0 and n_ag == 0:
        if n_at >= 2 or "11 N°6 Irreprochable" in atenuantes:
            idx_max = max(0, idx_min - 1)
            idx_min = max(0, idx_min - 1)
            efecto = "Rebaja de un grado"
        else:
            idx_max = idx_min
            efecto = "Mínimum del grado"
    elif n_ag > 0 and n_at == 0:
        idx_min = idx_max
        efecto = "Máximum del grado"
    elif n_at > 0 and n_ag > 0:
        efecto = "Compensación Racional (Rango completo)"
    else:
        efecto = "Sin modificatorias (Rango completo)"

    if es_rpa:
        idx_min = max(0, idx_min - 1)
        idx_max = max(0, idx_max - 1)
        efecto += " + Rebaja RPA Art. 21"

    rango_final = f"{ESCALA_PENAS[idx_min]['nombre']} a {ESCALA_PENAS[idx_max]['nombre']}"
    dias_min = ESCALA_PENAS[idx_min]['min']
    
    if es_rpa:
        if dias_min > 1825:
            resultado = "Régimen Cerrado (Crimen)"
            riesgo = 90
            badge = "badge-danger"
        elif dias_min > 1095:
            resultado = "Régimen Semicerrado"
            riesgo = 60
            badge = "badge-warning"
        else:
            resultado = "Libertad Asistida / Especial"
            riesgo = 20
            badge = "badge-success"
    else:
        if dias_min <= 1095:
            resultado = "Remisión Condicional (Probable)"
            riesgo = 10
            badge = "badge-success"
        elif dias_min <= 1825:
            resultado = "Libertad Vigilada (Probable)"
            riesgo = 40
            badge = "badge-warning"
        else:
            resultado = "Cumplimiento Efectivo"
            riesgo = 95
            badge = "badge-danger"

    return {
        "rango": rango_final,
        "dias_min": dias_min,
        "efecto": efecto,
        "resultado": resultado,
        "riesgo": riesgo,
        "badge": badge
    }

def generar_teoria_caso_ia(hechos, delito, atenuantes, es_rpa):
    contexto = "Adolescente (Ley 20.084)" if es_rpa else "Adulto"
    prompt = f"""
    Actúa como abogado penalista experto en litigación oral.
    Genera una TEORÍA DEL CASO estructurada para la defensa.
    DATOS DEL CASO:
    - Delito: {delito}
    - Contexto: {contexto}
    - Atenuantes invocadas: {", ".join(atenuantes)}
    - Relato de Hechos (Fiscalía): {hechos}
    ESTRUCTURA DE RESPUESTA REQUERIDA (NO USES MARKDOWN PESADO, SOLO TEXTO LIMPIO):
    1. PROPOSICIÓN FÁCTICA (Nuestra versión de los hechos, minimizando dolo o participación).
    2. PROPOSICIÓN JURÍDICA (Argumentos de derecho, calificación jurídica, improcedencia de prisión preventiva).
    3. PROPOSICIÓN PROBATORIA (Diligencias sugeridas: peritajes, testigos, documentos a solicitar).
    """
    try:
        response = model_ia.generate_content(prompt)
        return response.text
    except:
        return "Error conectando con IA Jurídica. Verifique conexión."

# =============================================================================
# 9. APLICACIÓN PRINCIPAL
# =============================================================================
def main_app():
    init_session_data()
    
    with st.sidebar:
        st.markdown(f"### 👤 {st.session_state.defensor_nombre}")
        if st.button("Cerrar Sesión"):
            st.session_state.logged_in = False
            st.rerun()
        st.divider()
        st.header("Gestión de Escritos")
        tipo_recurso = st.selectbox("Seleccionar Trámite", TIPOS_RECURSOS)
        st.divider()
        st.success("Supabase: Conectado (Ready)")
        st.info("BD Usuarios: Activa")
        st.info("BD Escritos: Activa")

    st.title(f"📄 {tipo_recurso}")
    
    tabs = st.tabs([
        "📝 Generador", 
        "🕵️ Analista Multimodal", 
        "🎙️ Transcriptor", 
        "📚 Biblioteca Inteligente", 
        "⚙️ Admin & BD"
    ])

    # === TAB 1: GENERADOR ===
    with tabs[0]:
        st.markdown("### 1. Individualización")
        
        # Botón de Limpieza General
        if st.button("🧼 Limpiar Todos los Campos", type="secondary"):
            st.session_state.defensor_nombre = ""
            st.session_state.imputado = ""
            st.session_state.lista_individualizacion = []
            st.rerun()

        col_def, col_imp = st.columns(2)
        st.session_state.defensor_nombre = col_def.text_input("Defensor/a", value=st.session_state.defensor_nombre)
        st.session_state.imputado = col_imp.text_input("Imputado/a", value=st.session_state.imputado)
        
        st.markdown("**Causas Individualizadas:**")
        for i, c in enumerate(st.session_state.lista_individualizacion):
            c1, c2, c3 = st.columns([3, 3, 1])
            c['rit'] = c1.text_input(f"RIT {i+1}", c['rit'], key=f"rit_ind_{i}")
            c['ruc'] = c2.text_input(f"RUC {i+1}", c['ruc'], key=f"ruc_ind_{i}")
            if c3.button("🗑️ Quitar", key=f"del_ind_{i}"):
                st.session_state.lista_individualizacion.pop(i)
                st.rerun()
                
        if st.button("➕ Agregar Causa a Individualización"):
            st.session_state.lista_individualizacion.append({"rit": "", "ruc": ""})
            st.rerun()
        
        tribunal_global = st.selectbox("Tribunal de Presentación", TRIBUNALES, index=TRIBUNALES.index(st.session_state.tribunal_sel) if st.session_state.tribunal_sel in TRIBUNALES else 0)
        st.session_state.tribunal_sel = tribunal_global

        st.markdown("---")
        
        if tipo_recurso == "Prescripción de la Pena":
            st.subheader("2. Causas a Prescribir (Detalle)")
            with st.form("form_prescripcion"):
                c1, c2, c3 = st.columns(3)
                p_rit = c1.text_input("RIT")
                p_ruc = c2.text_input("RUC")
                p_trib = c3.selectbox("Tribunal Origen", TRIBUNALES)
                c4, c5, c6 = st.columns(3)
                p_fecha_sent = c4.text_input("Fecha Sentencia", placeholder="12-12-2010")
                p_pena = c5.text_input("Pena Impuesta")
                p_delito = c6.text_input("Delito")
                p_fecha_susp = st.text_input("Fecha Ejecutoria / Suspensión")
                if st.form_submit_button("➕ Agregar Causa"):
                    st.session_state.prescripcion_list.append({
                        "rit": p_rit, "ruc": p_ruc, "tribunal": p_trib,
                        "fecha_sentencia": p_fecha_sent, "pena": p_pena,
                        "delito": p_delito, "fecha_suspension": p_fecha_susp
                    })
                    st.success("Causa agregada.")
            
            if st.session_state.prescripcion_list:
                st.write("**Causas en el escrito:**")
                for i, c in enumerate(st.session_state.prescripcion_list):
                    c1, c2 = st.columns([8, 1])
                    c1.caption(f"{i+1}. {c['delito']} (RIT {c['rit']})")
                    if c2.button("🗑️", key=f"del_pres_{i}"):
                        st.session_state.prescripcion_list.pop(i)
                        st.rerun()

        elif tipo_recurso == "Extinción Art. 25 ter":
            c_rpa, c_ad = st.columns(2)
            with c_rpa:
                st.markdown("#### A. Causa RPA")
                for i, rpa in enumerate(st.session_state.rpa):
                    with st.expander(f"Causa RPA {i+1}", expanded=True):
                        rpa['rit'] = st.text_input("RIT", rpa.get('rit',''), key=f"rrit{i}")
                        rpa['ruc'] = st.text_input("RUC", rpa.get('ruc',''), key=f"rruc{i}")
                        rpa['tribunal'] = st.selectbox("Tribunal", TRIBUNALES, key=f"rtrib{i}")
                        rpa['sancion'] = st.text_input("Sanción", rpa.get('sancion',''), key=f"rsanc{i}")
                        if st.button("🗑️ Quitar", key=f"del_rpa_{i}"):
                            st.session_state.rpa.pop(i)
                            st.rerun()
                if st.button("➕ Otra RPA"):
                    st.session_state.rpa.append({})
                    st.rerun()

            with c_ad:
                st.markdown("#### B. Condena Adulto")
                for i, ad in enumerate(st.session_state.adulto):
                    with st.expander(f"Condena Adulto {i+1}", expanded=True):
                        ad['rit'] = st.text_input("RIT", ad.get('rit',''), key=f"arit{i}")
                        ad['ruc'] = st.text_input("RUC", ad.get('ruc',''), key=f"aruc{i}")
                        ad['tribunal'] = st.selectbox("Tribunal", TRIBUNALES, key=f"atrib{i}")
                        ad['pena'] = st.text_input("Pena", ad.get('pena',''), key=f"apena{i}")
                        ad['fecha'] = st.text_input("Fecha", ad.get('fecha',''), key=f"afecha{i}")
                        if st.button("🗑️ Quitar", key=f"del_ad_{i}"):
                            st.session_state.adulto.pop(i)
                            st.rerun()
                if st.button("➕ Otra Adulto"):
                    st.session_state.adulto.append({})
                    st.rerun()

        elif tipo_recurso == "Apelación por Quebrantamiento":
            st.subheader("2. Detalle del Quebrantamiento")
            
            # Campos Específicos para Apelación
            col_ap1, col_ap2 = st.columns(2)
            rit_ap = col_ap1.text_input("RIT Causa Apelación")
            ruc_ap = col_ap2.text_input("RUC Causa Apelación")
            
            # CAMBIO: Agregado campo HECHOS
            hechos_quebrantamiento = st.text_area("Hechos del Quebrantamiento", height=100, placeholder="Describa brevemente qué ocurrió...")
            
            resolucion_tribunal = st.text_area("Resolución del Tribunal (Que se impugna)", height=100)
            
            # CAMBIO: Refuerzo IA para Argumentos
            argumentos_defensa = st.text_area("Argumentos Defensa (Borrador)", height=100)
            if st.button("✨ Robustecer Argumentos con IA"):
                with st.spinner("Mejorando argumentación jurídica..."):
                    try:
                        model_ia = get_generative_model_dinamico()
                        prompt_arg = f"""
                        Actúa como Abogado Defensor Penal experto.
                        Mejora y robustece estos argumentos para una Apelación por Quebrantamiento:
                        "{argumentos_defensa}"
                        
                        Usa terminología jurídica precisa, cita principios (proporcionalidad, interés superior adolescente) y mantén un tono persuasivo pero formal.
                        Solo entrega el texto mejorado.
                        """
                        resp_arg = model_ia.generate_content(prompt_arg)
                        argumentos_defensa = resp_arg.text
                        st.success("Argumentos mejorados. Copia y pega si es necesario.")
                        st.text_area("Argumentos Mejorados (Copia de aquí)", value=argumentos_defensa, height=150)
                    except Exception as e:
                        st.error(f"Error IA: {e}")

            antecedentes_sociales = st.text_area("Antecedentes Sociales (Opcional)", height=80, placeholder="Educacional, Laboral, Familiar...")
            
            col_san1, col_san2 = st.columns(2)
            sancion_orig = col_san1.text_input("Sanción Original")
            sancion_queb = col_san2.text_input("Sanción Quebrantada")
            
            st.session_state.datos_apelacion = {
                "rit_ap": rit_ap, "ruc_ap": ruc_ap,
                "hechos_quebrantamiento": hechos_quebrantamiento, # Nuevo campo
                "resolucion_tribunal": resolucion_tribunal,
                "argumentos_defensa": argumentos_defensa,
                "antecedentes_sociales": antecedentes_sociales,
                "sancion_orig": sancion_orig,
                "sancion_quebrantada": sancion_queb
            }

        elif tipo_recurso == "Amparo Constitucional":
            st.subheader("2. Fundamentos Específicos")
            argumento_extra = st.text_area("Antecedentes de Hecho Adicionales (Opcional)", height=150)
            st.session_state.argumento_extra = argumento_extra

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button(f"🚀 GENERAR ESCRITO: {tipo_recurso}", type="primary", use_container_width=True):
            dm_safe = st.session_state.get('datos_minuta', {})
            datos_apelacion = st.session_state.get('datos_apelacion', {})
            
            datos_finales = {
                "tribunal_ej": st.session_state.tribunal_sel,
                "prescripcion_list": st.session_state.prescripcion_list,
                "rpa": st.session_state.rpa,
                "adulto": st.session_state.adulto,
                "ejecucion": st.session_state.ejecucion,
                "lista_individualizacion": st.session_state.lista_individualizacion,
                "argumento_extra": st.session_state.get('argumento_extra', ''),
                "fecha_det": dm_safe.get('fecha', ''),
                "lugar_det": dm_safe.get('lugar', ''),
                "argumentos_det": dm_safe.get('args', []),
                "hechos_relato": dm_safe.get('hechos_relato', ''),
                "version_imputado": dm_safe.get('version_imputado', ''),
                # Campos Apelación
                "rit_ap": datos_apelacion.get('rit_ap', ''),
                "ruc_ap": datos_apelacion.get('ruc_ap', ''),
                "hechos_quebrantamiento": datos_apelacion.get('hechos_quebrantamiento', ''), # Nuevo
                "resolucion_tribunal": datos_apelacion.get('resolucion_tribunal', ''),
                "argumentos_defensa": datos_apelacion.get('argumentos_defensa', ''),
                "antecedentes_sociales": datos_apelacion.get('antecedentes_sociales', ''),
                "sancion_orig": datos_apelacion.get('sancion_orig', ''),
                "sancion_quebrantada": datos_apelacion.get('sancion_quebrantada', '')
            }
            gen = GeneradorWord(st.session_state.defensor_nombre, st.session_state.imputado)
            buffer = gen.generar(tipo_recurso, datos_finales)
            st.success("Documento Generado Exitosamente")
            st.download_button("📥 Descargar DOCX", buffer, f"{tipo_recurso}.docx", 
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                             use_container_width=True)

                     # --- ÁREA DE PROCESAMIENTO REAL (MASTER RPA) ---
        st.markdown("---")
        with st.expander("🛠️ PANEL DE CONTROL RPA & ANÁLISIS MAESTRO", expanded=True):
            st.markdown("### 🤖 Procesamiento Inteligente de la Causa")
            st.info("Este panel utiliza IA para analizar la totalidad del texto extraído de los documentos cargados en el sistema.")
            
            if st.button("🚀 INICIAR ANÁLISIS ESTRATÉGICO RPA", use_container_width=True):
                # Verificamos que exista texto acumulado en el estado de la sesión
                # Nota: 'all_text' se alimenta desde la pestaña de Analista o Ingesta
                if 'all_text' not in st.session_state or not st.session_state.all_text.strip():
                    st.warning("⚠️ No hay texto extraído para procesar. Por favor, sube archivos en 'Analista' o 'Admin' primero.")
                else:
                    with st.spinner("⚖️ Analizando documentos con Inteligencia Artificial..."):
                        try:
                            # 1. Definimos una consulta técnica de alto nivel
                            query_rpa = """
                            Realiza un análisis jurídico profundo de los documentos proporcionados.
                            TU TAREA:
                            1. RESUMEN EJECUTIVO: Puntos clave de las causas.
                            2. DETECCIÓN DE RIESGOS: Identifica plazos por vencer o debilidades procesales.
                            3. RECOMENDACIÓN: Sugiere la mejor vía de extinción o recurso aplicable (Estilo Defensoría Chile).
                            """
                            
                            # 2. Ejecutamos la llamada real al modelo configurado
                            resultado_ia = process_legal_query(query_rpa, st.session_state.all_text)
                            
                            # 3. Despliegue de resultados
                            st.markdown("---")
                            st.markdown("#### 📋 Informe de Análisis Legal")
                            st.markdown(resultado_ia)
                            
                            # 4. Registro en el log
                            timestamp = datetime.now().strftime("%H:%M:%S")
                            if 'logs' not in st.session_state: st.session_state.logs = []
                            st.session_state.logs.append(f"[{timestamp}] Análisis RPA masivo completado.")
                            st.success("✅ Procesamiento finalizado con éxito.")
                        
                        # === AQUÍ ESTABA EL ERROR: FALTABA ESTE BLOQUE EXCEPT ===
                        except Exception as e:
                            st.error(f"Error crítico en el motor de IA: {e}")


                            
    # === TAB 2: ANALISTA MULTIMODAL (MERGED FUNCTIONS + SUMMARY BOX) ===
    with tabs[1]:
        st.header("🕵️ Analista Jurídico Multimodal (Vision & Strategy)")
        st.info("Sube Carpetas Investigativas, Partes Policiales Escaneados, Fotos de Evidencia o Textos.")

        objetivo_analisis = st.radio(
            "¿Qué buscas en estos documentos?",
            ["📄 Control de Detención (Busca ilegalidades)", 
             "⚖️ Estrategia Integral (Teoría del Caso, Salidas & Prognosis)"],
            horizontal=True
        )

        archivos_evidencia = st.file_uploader(
            "Cargar Evidencia (PDF, JPG, PNG, TXT)", 
            type=["pdf", "jpg", "png", "txt", "jpeg"], 
            accept_multiple_files=True
        )

        contexto_usuario = st.text_area("Contexto adicional (Ej: 'El cliente dice que Carabineros mintió...')")

        if archivos_evidencia and st.button("⚡ ANALIZAR EVIDENCIA CON IA"):
            status_box = st.empty()
            with st.spinner("Procesando evidencia multimodal (Vision IA)..."):
                try:
                    model_analista = get_generative_model_dinamico()
                    docs_para_gemini = []
                    
                    for archivo in archivos_evidencia:
                        status_box.info(f"Subiendo a Gemini Vision: {archivo.name}...")
                        suffix = f".{archivo.name.split('.')[-1]}"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(archivo.getvalue())
                            tmp_path = tmp.name

                        f_gemini = genai.upload_file(tmp_path)
                        while f_gemini.state.name == "PROCESSING":
                            time.sleep(1)
                            f_gemini = genai.get_file(f_gemini.name)
                            
                        docs_para_gemini.append(f_gemini)
                        os.remove(tmp_path)

                    status_box.info("🧠 Generando estrategia jurídica integral...")

                    prompt_system = """
                    Eres un Estratega de Defensa Penal.
                    IMPORTANTE: Tu respuesta es para un abogado. NO incluyas código python, ni json raw, ni expliques que eres una IA.
                    Solo entrega el informe jurídico profesional.
                    """

                    if "Control de Detención" in objetivo_analisis:
                        prompt_especifico = """
                        TU MISIÓN: Detectar vicios de legalidad para un Control de Detención.
                        Genera también un RECUADRO DE RESUMEN al final con:
                        - Ilegalidad detectada: (Sí/No)
                        - Probabilidad de éxito: (Alta/Media/Baja)
                        - Argumento clave.
                        """
                    else:
                        prompt_especifico = """
                        TU MISIÓN: Construir una Estrategia de Defensa Integral.
                        
                        ESTRUCTURA OBLIGATORIA DEL INFORME:
                        1. ANÁLISIS DE LA PRUEBA (Debilidades fiscalía).
                        2. TEORÍA DEL CASO (Nuestra versión).
                        
                        AL FINAL, GENERA UN BLOQUE LLAMADO "RESUMEN ESTRATÉGICO" CON:
                        - Pena Probable: (Ej: 541 días)
                        - Pena Sustitutiva: (Ej: Remisión Condicional)
                        - Atenuantes: (Lista)
                        - Agravantes: (Lista)
                        - Salida Alternativa: (Viabilidad SCP o AR)
                        - Recomendación: (Juicio o Abreviado)
                        """

                    prompt_final = [prompt_system + prompt_especifico, f"Contexto adicional: {contexto_usuario}"]
                    prompt_final.extend(docs_para_gemini)

                    response = model_analista.generate_content(prompt_final)
                    
                    status_box.success("✅ Análisis Completado")
                    
               # Si tenías lógica residual aquí, nos aseguramos de cerrar el bloque try
                        # Para el Panel RPA, esto debería ser el final del éxito:
                        st.session_state.logs.append(f"[{datetime.now().strftime('%H:%M:%S')}] Operación completada.")
                        st.success("✅ Procesamiento finalizado con éxito.")

                    # === CIERRE OBLIGATORIO DEL BLOQUE TRY (CORRECCIÓN ERROR) ===
                    except Exception as e:
                        st.error(f"Error durante el procesamiento: {e}")

    # -----------------------------------------------------------------------------
    # === TAB 2: ANALISTA MULTIMODAL (MERGED FUNCTIONS + SUMMARY BOX) ===
    # -----------------------------------------------------------------------------
    with tabs[1]:
        st.header("🕵️ Analista Jurídico Multimodal (Vision & Strategy)")
        st.info("Sube Carpetas Investigativas, Partes Policiales Escaneados, Fotos de Evidencia o Textos.")

        objetivo_analisis = st.radio(
            "¿Qué buscas en estos documentos?",
            ["📄 Control de Detención (Busca ilegalidades)", 
             "⚖️ Estrategia Integral (Teoría del Caso, Salidas & Prognosis)"],
            horizontal=True
        )

        archivos_evidencia = st.file_uploader(
            "Cargar Evidencia (PDF, JPG, PNG, TXT)", 
            type=["pdf", "jpg", "png", "txt", "jpeg"], 
            accept_multiple_files=True
        )

        contexto_usuario = st.text_area("Contexto adicional (Ej: 'El cliente dice que Carabineros mintió...')")

        if archivos_evidencia and st.button("⚡ ANALIZAR EVIDENCIA CON IA"):
            status_box = st.empty()
            with st.spinner("Procesando evidencia multimodal (Vision IA)..."):
                try:
                    model_analista = get_generative_model_dinamico()
                    docs_para_gemini = []
                    
                    for archivo in archivos_evidencia:
                        status_box.info(f"Subiendo a Gemini Vision: {archivo.name}...")
                        suffix = f".{archivo.name.split('.')[-1]}"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(archivo.getvalue())
                            tmp_path = tmp.name

                        f_gemini = genai.upload_file(tmp_path)
                        while f_gemini.state.name == "PROCESSING":
                            time.sleep(1)
                            f_gemini = genai.get_file(f_gemini.name)
                            
                        docs_para_gemini.append(f_gemini)
                        os.remove(tmp_path)

                    status_box.info("🧠 Generando estrategia jurídica integral...")

                    prompt_system = """
                    Eres un Estratega de Defensa Penal.
                    IMPORTANTE: Tu respuesta es para un abogado. NO incluyas código python, ni json raw, ni expliques que eres una IA.
                    Solo entrega el informe jurídico profesional.
                    """

                    if "Control de Detención" in objetivo_analisis:
                        prompt_especifico = """
                        TU MISIÓN: Detectar vicios de legalidad para un Control de Detención.
                        Genera también un RECUADRO DE RESUMEN al final con:
                        - Ilegalidad detectada: (Sí/No)
                        - Probabilidad de éxito: (Alta/Media/Baja)
                        - Argumento clave.
                        """
                    else:
                        prompt_especifico = """
                        TU MISIÓN: Construir una Estrategia de Defensa Integral.
                        
                        ESTRUCTURA OBLIGATORIA DEL INFORME:
                        1. ANÁLISIS DE LA PRUEBA (Debilidades fiscalía).
                        2. TEORÍA DEL CASO (Nuestra versión).
                        
                        AL FINAL, GENERA UN BLOQUE LLAMADO "RESUMEN ESTRATÉGICO" CON:
                        - Pena Probable: (Ej: 541 días)
                        - Pena Sustitutiva: (Ej: Remisión Condicional)
                        - Atenuantes: (Lista)
                        - Agravantes: (Lista)
                        - Salida Alternativa: (Viabilidad SCP o AR)
                        - Recomendación: (Juicio o Abreviado)
                        """

                    prompt_final = [prompt_system + prompt_especifico, f"Contexto adicional: {contexto_usuario}"]
                    prompt_final.extend(docs_para_gemini)

                    response = model_analista.generate_content(prompt_final)
                    
                    status_box.success("✅ Análisis Completado")
                    
                    texto_resultado = response.text
                    
                    # Extracción simple del Resumen para mostrar en recuadro bonito
                    if "RESUMEN ESTRATÉGICO" in texto_resultado:
                        partes = texto_resultado.split("RESUMEN ESTRATÉGICO")
                        resumen_texto = partes[-1]
                        contenido_principal = partes[0]
                        st.markdown(f"<div class='resumen-dinamico'><h4>📊 RESUMEN ESTRATÉGICO</h4>{resumen_texto}</div>", unsafe_allow_html=True)
                        st.markdown(contenido_principal)
                    else:
                        st.markdown(texto_resultado)
                    
                    st.download_button("📥 Descargar Informe", texto_resultado, "Analisis_Integral_Legal.txt")

                except Exception as e:
                    st.error(f"Error en el análisis multimodal: {e}")

    # === TAB 3: TRANSCRIPTOR ===
    with tabs[2]:
        st.header("🎙️ Transcriptor Forense & Generador de Recursos")
        st.info("Sube el audio de la audiencia (MP3, WAV, M4A) para obtener la transcripción literal y un borrador de recurso inteligente.")

        uploaded_audio = st.file_uploader("Cargar Audio de Audiencia", type=["mp3", "wav", "m4a", "ogg"])

        if uploaded_audio is not None:
            if st.button("🚀 PROCESAR AUDIO (AUTO-DETECTAR MODELO)"):
                status_container = st.empty()
                with st.spinner("🔄 Auto-detectando modelo y procesando..."):
                    try:
                        model_transcriptor = get_generative_model_dinamico() # Usamos el getter dinámico
                        status_container.info(f"🤖 Procesando audio...")

                        suffix = f".{uploaded_audio.name.split('.')[-1]}"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                            tmp_file.write(uploaded_audio.getvalue())
                            tmp_path = tmp_file.name

                        archivo_gemini = genai.upload_file(tmp_path, mime_type="audio/mp3")

                        status_container.info("⏳ Esperando procesamiento de Google...")
                        while archivo_gemini.state.name == "PROCESSING":
                            time.sleep(2)
                            archivo_gemini = genai.get_file(archivo_gemini.name)

                        if archivo_gemini.state.name == "FAILED":
                            raise ValueError("Google falló al procesar el audio.")

                        status_container.info("📝 Redactando recurso...")
                        
                        prompt_transcripcion = """
                        Actúa como un Estenógrafo Judicial y Abogado Penalista.
                        TAREA 1: Transcribe LITERALMENTE el audio (Juez, Fiscal, Defensa).
                        TAREA 2: Redacta un BORRADOR DE RECURSO (Apelación o Amparo) detectando los vicios en el audio.
                        Estructura: Resolución Impugnada, Argumentos Defensa, Agravio, Petitorio.
                        """

                        response = model_transcriptor.generate_content([prompt_transcripcion, archivo_gemini])
                        texto_generado = response.text

                        status_container.success("✅ ¡Listo!")
                        st.subheader(f"📄 Resultado")
                        st.markdown(texto_generado)

                        st.download_button("📥 Descargar", texto_generado, "Recurso_Audiencia.txt")

                    except Exception as e:
                        st.error(f"Error: {e}")
                    finally:
                        if 'tmp_path' in locals() and os.path.exists(tmp_path):
                            os.remove(tmp_path)
        else:
            st.warning("Por favor, carga un archivo de audio para comenzar.")

    # === TAB 4: BIBLIOTECA INTELIGENTE (RAG MEJORADO + ANALISIS) ===
    with tabs[3]:
        st.header("📚 Biblioteca Jurídica Inteligente (Investigación RAG)")
        
        modo_biblio = st.radio("Herramienta", ["🔍 Buscador Jurídico Avanzado", "📄 Analizar mi Escrito"], horizontal=True)
        
        if modo_biblio == "🔍 Buscador Jurídico Avanzado":
            st.info("Buscador semántico potenciado por IA: Filtra, encuentra similitudes y genera respuestas jurídicas.")
            
            # Layout de Filtros
            col_filtros = st.columns(3)
            filtro_tipo = col_filtros[0].selectbox("Tipo de Documento", ["Todos", "Sentencia Condenatoria", "Sentencia Absolutoria", "Recurso de Nulidad", "Recurso de Amparo", "Recurso de Apelación", "Doctrina/Artículo", "Ley/Normativa"])
            filtro_tribunal = col_filtros[1].text_input("Tribunal (Opcional)", placeholder="Ej: Suprema, San Miguel")
            query_busqueda = col_filtros[2].text_input("Tema Jurídico / Consulta", placeholder="Ej: Nulidad por falta de fundamentación")
            
            if st.button("🔎 Investigar"):
                with st.spinner("Consultando bases de datos y generando respuesta..."):
                    try:
                        # 1. Obtener Embedding de la consulta
                        modelo_dinamico = get_embedding_model()
                        emb_resp = genai.embed_content(
                            model=modelo_dinamico,
                            content=query_busqueda,
                            task_type="retrieval_query"
                        )
                        vector_consulta = emb_resp['embedding']
                        
                        if vector_consulta:
                            # 2. Construir Query a Supabase con filtros
                            query_db = supabase.table("documentos_legales").select("*").limit(100)
                            
                            if filtro_tipo != "Todos":
                                query_db = query_db.eq('metadata->>tipo', filtro_tipo)
                            # Nota: El filtro de tribunal es texto parcial, mejor hacerlo en Python si no es exacto
                            
                            res = query_db.execute()
                            
                            if res.data:
                                resultados_scores = []
                                for doc in res.data:
                                    vec_doc = doc.get('embedding')
                                    # CORRECCIÓN ERROR TIPOS
                                    if isinstance(vec_doc, str):
                                        vec_doc = json.loads(vec_doc)
                                    
                                    # Filtro tribunal parcial en Python
                                    meta = doc['metadata']
                                    if isinstance(meta, str): meta = json.loads(meta)
                                    
                                    if filtro_tribunal and filtro_tribunal.lower() not in meta.get('tribunal', '').lower():
                                        continue

                                    if vec_doc:
                                        v_a = np.array(vector_consulta)
                                        v_b = np.array(vec_doc)
                                        sim = np.dot(v_a, v_b) / (np.linalg.norm(v_a) * np.linalg.norm(v_b))
                                        resultados_scores.append((sim, doc, meta))
                                
                                # Ordenar top 5
                                resultados_scores.sort(key=lambda x: x[0], reverse=True)
                                top_resultados = resultados_scores[:5]
                                
                                if top_resultados:
                                    # 3. GENERACIÓN DE RESPUESTA JURÍDICA (RAG)
                                    contexto_rag = ""
                                    for i, (score, doc, meta) in enumerate(top_resultados):
                                        contexto_rag += f"FRAGMENTO {i+1} (Rol: {meta.get('rol')}, Tribunal: {meta.get('tribunal')}): {doc['contenido'][:800]}...\n\n"
                                    
                                    prompt_rag = f"""
                                    Eres un abogado investigador senior.
                                    Basado EXCLUSIVAMENTE en estos fragmentos de jurisprudencia recuperados:
                                    {contexto_rag}
                                    
                                    Redacta una respuesta jurídica directa a la consulta: '{query_busqueda}'.
                                    Cita obligatoriamente los ROLES y TRIBUNALES de cada fragmento usado para respaldar tu respuesta.
                                    Si la información no es suficiente, indícalo.
                                    """
                                    
                                    model_resp = get_generative_model_dinamico()
                                    resp_juridica = model_resp.generate_content(prompt_rag)
                                    
                                    st.markdown("<div class='resumen-dinamico'><h4>⚖️ RESPUESTA JURÍDICA INTELIGENTE</h4>" + resp_juridica.text + "</div>", unsafe_allow_html=True)
                                    
                                    st.divider()
                                    st.caption("FUENTES CONSULTADAS:")
                                    
                                    # 4. Mostrar Fuentes con Badges
                                    for score, doc, meta in top_resultados:
                                        tipo_doc = meta.get('tipo', 'Doc')
                                        rol_doc = meta.get('rol', 'S/N')
                                        trib_doc = meta.get('tribunal', '')
                                        
                                        with st.expander(f"{trib_doc} - {rol_doc} (Relevancia: {int(score*100)}%)"):
                                            st.markdown(f"<span class='badge-tipo'>{tipo_doc}</span> <span class='badge-rol'>{rol_doc}</span>", unsafe_allow_html=True)
                                            st.markdown(f"**Resultado:** {meta.get('resultado', '-')}")
                                            st.write(doc['contenido'][:1000] + "...")
                                            st.button("Copiar Cita", key=f"btn_{doc['id']}")
                                else:
                                    st.warning("No se encontraron coincidencias relevantes con esos filtros.")
                            else:
                                st.warning("La base de datos no tiene documentos que coincidan con el filtro inicial.")
                        else:
                            st.error("Error generando vector de búsqueda.")

                    except Exception as e:
                        st.error(f"Error en motor de búsqueda: {e}")

        else: # Analizar mi Escrito (MEJORADO: SUGERENCIAS DIRECTAS)
            st.info("Sube tu borrador. La IA detectará debilidades y sugerirá argumentos de derecho sólidos.")
            borrador = st.file_uploader("Sube tu borrador (PDF/Word/Txt)", type=["pdf","docx","txt"])
            
           # --- ANÁLISIS REAL DE ESTRATEGIA JURÍDICA ---
            st.info("💡 Análisis estratégico real mediante LangChain y Gemini 1.5 Pro.")
            
            if borrador and st.button("⚖️ Ejecutar Análisis Estratégico Real"):
                with st.spinner("Analizando estrategia jurídica con IA..."):
                    try:
                        # Extraemos el texto del borrador subido
                        texto_borrador = extraer_texto_generico(borrador)
                        
                        if texto_borrador:
                            prompt_analisis_escrito = """
                            Actúa como un Abogado Senior y Profesor de Derecho Penal. Analiza el borrador adjunto.
                            NO RESUMAS EL DOCUMENTO. VE DIRECTO AL GRANO.
                            
                            TU TAREA ES ENTREGAR:
                            1. 🚩 DEBILIDADES DETECTADAS: ¿Qué argumento es débil o falta fundamentación?
                            2. 🛡️ SUGERENCIAS DE DERECHO: Redacta párrafos jurídicos sólidos para reforzar esas debilidades.
                            3. ⚖️ JURISPRUDENCIA SUGERIDA: Indica líneas de fallos específicas a buscar.
                            """
                            
                            # Llamada real a la IA
                            respuesta_real = process_legal_query(prompt_analisis_escrito, texto_borrador)
                            
                            st.success("✅ Análisis Estratégico Completado")
                            st.markdown("---")
                            st.markdown(respuesta_real)
                            st.session_state.logs.append(f"Análisis de borrador '{borrador.name}' completado.")
                        else:
                            st.error("No se pudo extraer texto del borrador.")
                            
                    except Exception as e:
                        st.error(f"Error en el análisis real: {e}")

    # === TAB 5: ADMIN & CARGA (GESTIÓN USUARIOS + INGESTA DINÁMICA + OCR) ===
    with tabs[4]:
        if st.session_state.user_role == "Admin":
            st.header("⚙️ Cerebro Centralizado & Gestión (Admin)")
            
            tab_ingesta, tab_usuarios = st.tabs(["📂 Ingesta Documental", "👥 Gestión de Usuarios"])
            
            # --- SUB-TAB A: INGESTA ---
            with tab_ingesta:
                st.info("Alimenta el sistema con Leyes y Jurisprudencia. Proceso inteligente con IA.")
                col_subida, col_consulta = st.columns([1, 1])

                with col_subida:
                    st.subheader("1. Ingesta Inteligente")
                    
                    archivos_pdf = st.file_uploader(
                        "Subir Archivos (PDF) - Máx 10", 
                        type="pdf", 
                        accept_multiple_files=True,
                        key="pdf_rag_multi"
                    )

                    if archivos_pdf:
                        if len(archivos_pdf) > 10:
                            st.error("⚠️ Por estabilidad y seguridad, sube máximo 10 archivos a la vez.")
                            st.stop()

                        if st.button("💾 Procesar y Guardar en Memoria"):
                            progress_bar_general = st.progress(0)
                            total_files = len(archivos_pdf)
                            
                            modelo_dinamico = get_embedding_model()
                            st.write(f"Usando modelo de embedding: {modelo_dinamico}")
                            
                            for idx_file, archivo_pdf in enumerate(archivos_pdf):
                                with st.status(f"Procesando {archivo_pdf.name}...", expanded=False) as status:
                                    try:
                                        status.write("Leyendo documento completo...")
                                        reader = PyPDF2.PdfReader(archivo_pdf)
                                        texto_completo = ""
                                        for page in reader.pages:
                                            texto_completo += page.extract_text() or ""
                                        
                                        # CAMBIO: OCR HÍBRIDO (Si hay poco texto, usamos Vision)
                                        if len(texto_completo) < 50:
                                            status.write("⚠️ Texto insuficiente, activando OCR con IA Vision...")
                                            st.toast(f"Activando OCR para {archivo_pdf.name}")
                                            
                                            suffix = f".{archivo_pdf.name.split('.')[-1]}"
                                            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                                                tmp.write(archivo_pdf.getvalue())
                                                tmp_path = tmp.name
                                            
                                            f_gemini = genai.upload_file(tmp_path)
                                            while f_gemini.state.name == "PROCESSING": 
                                                time.sleep(1)
                                                f_gemini = genai.get_file(f_gemini.name)
                                            
                                            model_ocr = get_generative_model_dinamico()
                                            prompt_ocr = """
                                            Analiza este documento legal escaneado.
                                            1. Extrae el TEXTO COMPLETO (transcripción literal).
                                            2. Genera un JSON con metadata: tribunal, rol, fecha_sentencia, resultado, tema, tipo.
                                            IMPORTANTE: El 'tipo' debe ser uno de: ["Sentencia Condenatoria", "Sentencia Absolutoria", "Recurso de Nulidad", "Recurso de Amparo", "Recurso de Apelación", "Doctrina/Artículo", "Ley/Normativa"].
                                            FORMATO RESPUESTA:
                                            ---JSON---
                                            {json_aqui}
                                            ---TEXTO---
                                            (texto_aqui)
                                            """
                                            resp_ocr = model_ocr.generate_content([prompt_ocr, f_gemini])
                                            
                                            parts = resp_ocr.text.split("---TEXTO---")
                                            json_part = parts[0].replace("---JSON---", "").replace("```json", "").replace("```", "").strip()
                                            texto_completo = parts[1].strip() if len(parts) > 1 else ""
                                            try:
                                                metadata_ia = json.loads(json_part)
                                            except:
                                                metadata_ia = {"rol": "Error OCR", "tribunal": "Desconocido", "tipo": "Documento Legal"}
                                            
                                            os.remove(tmp_path)

                                        else:
                                            status.write("Analizando metadata jurídica con IA...")
                                            metadata_ia = analizar_metadata_profunda(texto_completo)
                                        
                                        metadata_ia["origen"] = archivo_pdf.name
                                        status.write(f"Metadata detectada: {metadata_ia.get('rol')} - {metadata_ia.get('tribunal')}")

                                        status.write("Fragmentando texto...")
                                        chunk_size = 1500 
                                        chunks = [texto_completo[i:i+chunk_size] for i in range(0, len(texto_completo), chunk_size)]
                                        
                                        status.write("Generando vectores y guardando...")
                                        for i, chunk in enumerate(chunks):
                                            emb_resp = genai.embed_content(
                                                model=modelo_dinamico,
                                                content=chunk,
                                                task_type="retrieval_document"
                                            )
                                            vector = emb_resp['embedding']

                                            if vector:
                                                data_insert = {
                                                    "contenido": chunk,
                                                    "metadata": metadata_ia,
                                                    "embedding": vector
                                                }
                                                supabase.table("documentos_legales").insert(data_insert).execute()
                                        
                                        status.update(label=f"✅ {archivo_pdf.name} Procesado Exitosamente", state="complete")
                                        st.toast(f"✅ Guardado: {metadata_ia.get('rol')} - {metadata_ia.get('tribunal')}")

                                    except Exception as e:
                                        status.update(label=f"❌ Error en {archivo_pdf.name}: {str(e)}", state="error")
                                        st.error(f"Detalle error: {e}")
                                
                                progress_bar_general.progress((idx_file + 1) / total_files)

                            st.success("🏁 Proceso de ingesta finalizado.")
                            time.sleep(2)
                            st.rerun()

                with col_consulta:
                    st.subheader("2. Inventario Documental")
                    # LÓGICA DE INVENTARIO MEJORADA (SOLICITUD USUARIO)
                    try:
                        res = supabase.table("documentos_legales").select("metadata, id").order("id", desc=True).limit(20).execute()
                        
                        if res.data:
                            data_limpia = []
                            for d in res.data:
                                m = d.get('metadata', {})
                                if isinstance(m, str): 
                                    try: m = json.loads(m)
                                    except: m = {}
                                
                                data_limpia.append({
                                    "ID": d['id'],
                                    "Tribunal": m.get('tribunal', 'N/A'),
                                    "Rol": m.get('rol', 'S/N'),
                                    "Tipo": m.get('tipo', 'Doc')
                                })
                            
                            st.dataframe(data_limpia, use_container_width=True, hide_index=True)
                        else:
                            st.info("La base de datos está vacía.")
                            
                    except Exception as e:
                        st.error(f"Error cargando inventario: {e}")
            
            # --- SUB-TAB B: USUARIOS ---
            with tab_usuarios:
                st.subheader("👥 Gestión de Usuarios del Sistema")
                
                c_lista, c_crear = st.columns([2, 1])
                
                with c_lista:
                    st.markdown("##### Usuarios Registrados")
                    try:
                        users_data = supabase.table("profiles").select("*").execute()
                        if users_data.data:
                            clean_users = []
                            for u in users_data.data:
                                clean_users.append({
                                    "Nombre": u.get('nombre', 'Sin Nombre'),
                                    "Rol": u.get('rol', 'User'),
                                    "Fecha Registro": u.get('created_at', '')[:10]
                                })
                            st.dataframe(clean_users, use_container_width=True)
                        else:
                            st.info("No se encontraron perfiles de usuario.")
                    except Exception as e:
                        st.error(f"Error al cargar usuarios: {e}")

                with c_crear:
                    st.markdown("##### Registrar Nuevo Funcionario")
                    with st.form("admin_create_user"):
                        new_u_email = st.text_input("Correo Institucional")
                        new_u_pass = st.text_input("Contraseña Temporal", type="password")
                        new_u_name = st.text_input("Nombre Funcionario")
                        new_u_role = st.selectbox("Rol Asignado", ["User", "Admin"])
                        
                        btn_crear = st.form_submit_button("Crear Usuario")
                        
                        if btn_crear:
                            try:
                                res = supabase.auth.sign_up({
                                    "email": new_u_email,
                                    "password": new_u_pass,
                                    "options": {
                                        "data": {
                                            "nombre": new_u_name,
                                            "rol_solicitado": new_u_role 
                                        }
                                    }
                                })
                                
                                if res.user:
                                    time.sleep(1)
                                    supabase.table("profiles").update({"rol": new_u_role}).eq("id", res.user.id).execute()
                                    st.success(f"Usuario {new_u_name} creado correctamente.")
                                    st.warning("⚠️ Nota: Es posible que debas volver a iniciar sesión si el sistema te cambió de cuenta automáticamente.")
                                else:
                                    st.error("No se pudo crear el usuario. Verifique el correo.")
                                    
                            except Exception as e:
                                st.error(f"Error creando usuario: {e}")

        else:
            st.warning("🔒 Acceso restringido a Administradores.")
            st.info("Debes iniciar sesión con una cuenta autorizada.")

if __name__ == "__main__":
    if st.session_state.logged_in:
        main_app()
    else:
        login_screen()
