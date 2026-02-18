import reflex as rx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import io
import re
import json
import os
import tempfile
import time
import base64
import numpy as np
import PyPDF2
from supabase import create_client
import google.generativeai as genai
from datetime import datetime

# =============================================================================
# 1. CONFIGURACIÓN Y ESTILOS GLOBALES
# =============================================================================

# Configuración de Claves (Se recomienda usar variables de entorno en producción)
SUPABASE_URL = os.environ.get("SUPABASE_URL", "") 
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")

# Configuración de servicios
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)

# BLOQUE DE COLORES CORREGIDO
COLORS = {
    "primary": "#161B2F",    # Navy Profundo
    "secondary": "#5B687C",  # Slate Blue
    "accent": "#D4CDCB",     # Beige
    "background": "#F4F7F6", # Gris Suave
    "text": "#161B2F",       # Texto Oscuro
    "error": "#E53935",      # ROJO (ESTE FALTABA)
    "success": "#43A047"     # Verde (Por si acaso)
}

# Estilos de Componentes Reutilizables
style_card = {
    "bg": COLORS["white"],
    "border": f"1px solid {COLORS['beige']}",
    "border_radius": "12px",
    "padding": "1.5em",
    "box_shadow": "0 4px 6px rgba(0, 0, 0, 0.05)",
}

style_input = {
    "border": f"1px solid {COLORS['beige']}",
    "bg": "white",
    "border_radius": "6px",
    "padding": "0.5em",
    "width": "100%",
}

style_button_primary = {
    "bg": COLORS["navy"],
    "color": "white",
    "_hover": {"bg": "#2C3550"},
    "width": "100%",
    "padding": "1em",
    "border_radius": "8px",
}

# =============================================================================
# 2. LÓGICA DE SOPORTE (Helpers fuera del State)
# =============================================================================

def get_supabase():
    if SUPABASE_URL and SUPABASE_KEY:
        return create_client(SUPABASE_URL, SUPABASE_KEY)
    return None

def get_generative_model_dinamico():
    """Busca automáticamente un modelo generativo disponible."""
    try:
        modelos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        mejor = next((m for m in modelos if 'gemini-1.5-flash' in m), None)
        if not mejor:
            mejor = next((m for m in modelos if 'gemini-1.5-pro' in m), modelos[0])
        return genai.GenerativeModel(mejor)
    except:
        return genai.GenerativeModel('models/gemini-1.5-flash-latest')

def get_embedding_model():
    """Busca modelo de embedding disponible."""
    try:
        # Lógica simplificada para Reflex
        return 'models/text-embedding-004'
    except:
        return 'models/text-embedding-004'

class GeneradorWord:
    """Clase auxiliar para generar documentos Word (Migrada de app.py)"""
    def __init__(self, defensor, imputado):
        self.doc = Document()
        self.defensor = defensor.upper() if defensor else "DEFENSOR PÚBLICO"
        self.imputado = imputado.upper() if imputado else "IMPUTADO"
        
        section = self.doc.sections[0]
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(12)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_parrafo(self, texto, negrita=False, align="JUSTIFY", sangria=True):
        p = self.doc.add_paragraph()
        if align == "CENTER": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "LEFT": p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if sangria and align == "JUSTIFY":
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        texto_final = texto.replace("{DEFENSOR}", self.defensor).replace("{IMPUTADO}", self.imputado)
        
        if negrita:
            run = p.add_run(texto_final)
            run.font.name = 'Cambria'
            run.font.size = Pt(12)
            run.bold = True
        else:
            # Lógica simple de negrita para Reflex
            run = p.add_run(texto_final)
            run.font.name = 'Cambria'
            run.font.size = Pt(12)

    def generar(self, tipo, datos):
        # Lógica resumida de generación basada en el original
        sumas = {
            "Extinción Art. 25 ter": "EN LO PRINCIPAL: SOLICITA EXTINCIÓN; OTROSÍ: ACOMPAÑA DOCUMENTO.",
            "Prescripción de la Pena": "EN LO PRINCIPAL: Solicita Audiencia de Prescripción; OTROSÍ: Oficia.",
            "Amparo Constitucional": "EN LO PRINCIPAL: ACCIÓN DE AMPARO; OTROSÍ: ORDEN DE NO INNOVAR.",
            "Apelación por Quebrantamiento": "EN LO PRINCIPAL: APELACIÓN; OTROSÍ: NOTIFICACIÓN."
        }
        self.add_parrafo(sumas.get(tipo, "SOLICITUD"), negrita=True, align="LEFT", sangria=False)
        self.doc.add_paragraph() 
        
        destinatario = "ILTMA. CORTE DE APELACIONES" if "Apelación" in tipo or "Amparo" in tipo else datos.get('tribunal_ej', 'TRIBUNAL').upper()
        self.add_parrafo(destinatario, negrita=True, align="CENTER", sangria=False)
        self.doc.add_paragraph()

        intro = f"{{DEFENSOR}}, Abogada, Defensora Penal Pública, en representación de {{IMPUTADO}}, a S.S. respetuosamente digo:"
        self.add_parrafo(intro)

        # Cuerpo simplificado para el ejemplo (pero funcional)
        if tipo == "Prescripción de la Pena":
            self.add_parrafo("Que, por medio de la presente, vengo en solicitar a S.S. se sirva fijar día y hora para celebrar audiencia...")
            for c in datos.get('prescripcion_list', []):
                self.add_parrafo(f"En la causa RIT {c['rit']}: Condenado a {c['pena']}.")
        
        elif tipo == "Apelación por Quebrantamiento":
            self.add_parrafo("Que encontrándome dentro del plazo legal, vengo en interponer recurso de apelación...")
            self.add_parrafo("I. HECHOS:", negrita=True)
            self.add_parrafo(datos.get('hechos_quebrantamiento', ''))
            self.add_parrafo("ARGUMENTOS DE DERECHO:", negrita=True)
            self.add_parrafo(datos.get('argumentos_defensa', ''))

        self.add_parrafo("POR TANTO,", sangria=False)
        self.add_parrafo("SOLICITO A S.S. acceder a lo solicitado.", sangria=False)

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

# =============================================================================
# 3. ESTADO (STATE) - LÓGICA DE NEGOCIO Y DATOS
# =============================================================================

class State(rx.State):
    # --- Sesión ---
    logged_in: bool = False
    user_email: str = ""
    user_role: str = "user"
    user_name: str = ""
    current_page: str = "Generador"
    
    # Login Inputs
    login_email: str = ""
    login_pass: str = ""
    
    # --- Generador Vars ---
    defensor_nombre: str = ""
    imputado: str = ""
    tipo_recurso: str = "Prescripción de la Pena"
    tribunal_sel: str = "7° Juzgado de Garantía de Santiago"
    
    # Inputs dinámicos
    rit_input: str = ""
    ruc_input: str = ""
    pena_input: str = ""
    hechos_quebrantamiento: str = ""
    resolucion_tribunal: str = ""
    argumentos_defensa: str = ""
    antecedentes_sociales: str = ""
    
    prescripcion_list: list[dict] = []
    
    # --- Analista Vars ---
    analisis_objetivo: str = "Control de Detención"
    contexto_analisis: str = ""
    analisis_result: str = ""
    is_analyzing: bool = False
    
    # --- Biblioteca Vars ---
    busqueda_query: str = ""
    filtro_tipo: str = "Todos"
    filtro_tribunal: str = ""
    resultados_biblioteca: list[dict] = []
    respuesta_juridica_ia: str = ""
    is_searching: bool = False
    
    # --- Admin Vars ---
    admin_upload_status: str = ""
    is_ingesting: bool = False

    # --- FUNCIONES DE AUTH ---
    def login(self):
        sb = get_supabase()
        if not sb:
            return rx.window_alert("Error: Supabase no configurado.")
        
        try:
            session = sb.auth.sign_in_with_password({"email": self.login_email, "password": self.login_pass})
            user = session.user
            data = sb.table("profiles").select("*").eq("id", user.id).execute()
            
            if data.data:
                perfil = data.data[0]
                self.user_role = perfil.get('rol', 'user')
                self.user_name = perfil.get('nombre', 'Usuario')
                self.logged_in = True
            else:
                return rx.window_alert("Usuario sin perfil asignado.")
        except Exception as e:
            return rx.window_alert(f"Credenciales inválidas: {str(e)}")

    def logout(self):
        self.logged_in = False
        self.user_email = ""

    # --- FUNCIONES GENERADOR ---
    def add_prescripcion_item(self):
        self.prescripcion_list.append({
            "rit": self.rit_input,
            "ruc": self.ruc_input,
            "pena": self.pena_input
        })
        self.rit_input = ""
        self.ruc_input = ""
        self.pena_input = ""

    def clear_form(self):
        self.defensor_nombre = ""
        self.imputado = ""
        self.rit_input = ""
        self.prescripcion_list = []
        self.hechos_quebrantamiento = ""
        self.argumentos_defensa = ""

    def robustecer_argumentos(self):
        """Llama a Gemini para mejorar la argumentación"""
        if not self.argumentos_defensa or not GOOGLE_API_KEY:
            return
        
        try:
            model = get_generative_model_dinamico()
            prompt = f"Mejora estos argumentos jurídicos para una apelación penal: '{self.argumentos_defensa}'. Usa lenguaje técnico y formal."
            resp = model.generate_content(prompt)
            self.argumentos_defensa = resp.text
        except Exception as e:
            return rx.window_alert(f"Error IA: {str(e)}")

    def download_docx(self):
        """Genera el DOCX en memoria y dispara la descarga"""
        datos = {
            "tribunal_ej": self.tribunal_sel,
            "prescripcion_list": self.prescripcion_list,
            "hechos_quebrantamiento": self.hechos_quebrantamiento,
            "resolucion_tribunal": self.resolucion_tribunal,
            "argumentos_defensa": self.argumentos_defensa,
            "antecedentes_sociales": self.antecedentes_sociales,
            "rit_ap": self.rit_input,
            "ruc_ap": self.ruc_input
        }
        
        gen = GeneradorWord(self.defensor_nombre, self.imputado)
        buffer = gen.generar(self.tipo_recurso, datos)
        
        # En Reflex usamos base64 para descargar blobs generados en backend
        b64_data = base64.b64encode(buffer.getvalue()).decode()
        return rx.download(
            data=b64_data,
            filename=f"{self.tipo_recurso.replace(' ', '_')}.docx"
        )

    # --- FUNCIONES ANALISTA (Uploads) ---
    async def handle_analisis_upload(self, files: list[rx.UploadFile]):
        """Maneja subida de archivos y análisis con Gemini Vision"""
        self.is_analyzing = True
        self.analisis_result = "Procesando archivos..."
        yield
        
        if not GOOGLE_API_KEY:
            self.analisis_result = "Falta API Key."
            self.is_analyzing = False
            return

        docs_gemini = []
        model = get_generative_model_dinamico()

        try:
            for file in files:
                upload_data = await file.read()
                # Guardamos temporalmente para que la API de Gemini lo lea (requiere path)
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.filename.split('.')[-1]}") as tmp:
                    tmp.write(upload_data)
                    tmp_path = tmp.name
                
                f_gemini = genai.upload_file(tmp_path)
                while f_gemini.state.name == "PROCESSING":
                    time.sleep(1)
                    f_gemini = genai.get_file(f_gemini.name)
                
                docs_gemini.append(f_gemini)
                os.remove(tmp_path)

            prompt = f"Eres un Abogado Experto. Analiza la evidencia para: {self.analisis_objetivo}. Contexto: {self.contexto_analisis}"
            final_content = [prompt] + docs_gemini
            
            resp = model.generate_content(final_content)
            self.analisis_result = resp.text

        except Exception as e:
            self.analisis_result = f"Error: {str(e)}"
        
        self.is_analyzing = False

    # --- FUNCIONES BIBLIOTECA (RAG) ---
    def buscar_jurisprudencia(self):
        self.is_searching = True
        self.respuesta_juridica_ia = ""
        yield
        
        sb = get_supabase()
        if not sb:
            self.is_searching = False
            return

        try:
            # 1. Embedding Query
            emb_resp = genai.embed_content(
                model="models/text-embedding-004",
                content=self.busqueda_query,
                task_type="retrieval_query"
            )
            vector = emb_resp['embedding']
            
            # 2. Fetch & Filter (Híbrido)
            res = sb.table("documentos_legales").select("*").limit(50).execute()
            
            if res.data:
                resultados = []
                for doc in res.data:
                    meta = doc.get('metadata', {})
                    if isinstance(meta, str): meta = json.loads(meta)
                    
                    if self.filtro_tipo != "Todos" and meta.get('tipo') != self.filtro_tipo:
                        continue
                    
                    vec_doc = doc.get('embedding')
                    if isinstance(vec_doc, str): vec_doc = json.loads(vec_doc)
                    
                    if vec_doc and vector:
                        v_a = np.array(vector)
                        v_b = np.array(vec_doc)
                        sim = np.dot(v_a, v_b) / (np.linalg.norm(v_a) * np.linalg.norm(v_b))
                        doc['similarity'] = float(sim)
                        doc['meta_parsed'] = meta
                        resultados.append(doc)
                
                resultados.sort(key=lambda x: x['similarity'], reverse=True)
                self.resultados_biblioteca = resultados[:5]
                
                # 3. Generate Answer
                contexto = "\n".join([f"- {d['meta_parsed'].get('rol')}: {d['contenido'][:600]}" for d in self.resultados_biblioteca])
                model = get_generative_model_dinamico()
                prompt = f"Responde jurídicamente a '{self.busqueda_query}' usando SOLO estos fallos:\n{contexto}"
                resp = model.generate_content(prompt)
                self.respuesta_juridica_ia = resp.text
            else:
                self.resultados_biblioteca = []
                self.respuesta_juridica_ia = "No se encontraron resultados."

        except Exception as e:
            self.respuesta_juridica_ia = f"Error de búsqueda: {str(e)}"
        
        self.is_searching = False

    # --- FUNCIONES ADMIN (Ingesta) ---
    async def handle_ingesta_upload(self, files: list[rx.UploadFile]):
        if self.user_role != "Admin": return
        self.is_ingesting = True
        self.admin_upload_status = "Iniciando ingesta..."
        yield

        sb = get_supabase()
        docs_processed = 0
        
        try:
            for file in files:
                data = await file.read()
                # Proceso de lectura PDF (similar a app.py)
                reader = PyPDF2.PdfReader(io.BytesIO(data))
                text = "".join([p.extract_text() for p in reader.pages])
                
                # Si el texto es muy corto, usar OCR híbrido (Gemini Vision)
                if len(text) < 50:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(data)
                        tmp_path = tmp.name
                    
                    f_gemini = genai.upload_file(tmp_path)
                    while f_gemini.state.name == "PROCESSING": time.sleep(1); f_gemini = genai.get_file(f_gemini.name)
                    
                    model = get_generative_model_dinamico()
                    ocr_resp = model.generate_content(["Extrae todo el texto y metadata JSON de este legal.", f_gemini])
                    text = ocr_resp.text # Simplificación para el ejemplo
                    os.remove(tmp_path)

                # Metadata con IA
                meta_prompt = f"Extrae JSON metadata (rol, tribunal, tipo, tema) de: {text[:5000]}"
                model = get_generative_model_dinamico()
                meta_resp = model.generate_content(meta_prompt)
                try:
                    clean_json = meta_resp.text.replace('```json', '').replace('```', '').strip()
                    metadata = json.loads(clean_json)
                except:
                    metadata = {"rol": "Desconocido", "tipo": "Documento"}

                # Chunking & Embedding
                chunk_size = 1500
                chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
                
                for chunk in chunks:
                    emb_resp = genai.embed_content(model="models/text-embedding-004", content=chunk)
                    
                    sb.table("documentos_legales").insert({
                        "contenido": chunk,
                        "metadata": metadata,
                        "embedding": emb_resp['embedding']
                    }).execute()
                
                docs_processed += 1
            
            self.admin_upload_status = f"✅ Ingesta completa. {docs_processed} documentos procesados."

        except Exception as e:
            self.admin_upload_status = f"❌ Error: {str(e)}"
        
        self.is_ingesting = False


# =============================================================================
# 4. COMPONENTES DE INTERFAZ (UI)
# =============================================================================

def sidebar_item(text, icon, page):
    return rx.button(
        rx.hstack(
            rx.text(icon),
            rx.text(text, font_weight="500"),
            spacing="3",
        ),
        on_click=lambda: State.set_current_page(page),
        bg="transparent",
        color="white",
        _hover={"bg": COLORS['slate']},
        width="100%",
        justify_content="start",
        padding="1em"
    )

def sidebar():
    return rx.vstack(
        rx.heading("SISTEMA IABL", color="white", size="6", margin_bottom="2em"),
        sidebar_item("Generador", "📝", "Generador"),
        sidebar_item("Analista IA", "🧠", "Analista"),
        sidebar_item("Biblioteca", "📚", "Biblioteca"),
        sidebar_item("Admin", "👥", "Admin"),
        rx.spacer(),
        rx.box(
            rx.text(State.user_name, color=COLORS['beige'], size="2"),
            rx.button("Cerrar Sesión", on_click=State.logout, bg=COLORS['error'], width="100%", size="2"),
            padding="1em",
            width="100%"
        ),
        bg=COLORS['navy'],
        width="250px",
        height="100vh",
        padding="1.5em",
        position="sticky",
        top="0",
        left="0",
        align_items="start"
    )

def login_page():
    return rx.center(
        rx.vstack(
            rx.heading("SISTEMA JURÍDICO IABL", size="8", color=COLORS['navy'], text_align="center"),
            rx.text('"Automatización inteligente para defensores: tu tiempo vale."', 
                    color=COLORS['slate'], font_style="italic", margin_bottom="2em"),
            
            rx.card(
                rx.vstack(
                    rx.input(placeholder="Correo Institucional", on_change=State.set_login_email, style=style_input),
                    rx.input(placeholder="Contraseña", type="password", on_change=State.set_login_pass, style=style_input),
                    rx.button("INGRESAR AL SISTEMA", on_click=State.login, style=style_button_primary),
                    spacing="4"
                ),
                style=style_card,
                width="400px"
            ),
            
            rx.hstack(
                rx.vstack(rx.icon("file-text", size=30), rx.text("Redacción"), align_items="center"),
                rx.vstack(rx.icon("eye", size=30), rx.text("Visión IA"), align_items="center"),
                rx.vstack(rx.icon("library", size=30), rx.text("Biblioteca"), align_items="center"),
                spacing="6",
                margin_top="3em",
                color=COLORS['slate']
            ),
            align_items="center",
            spacing="2"
        ),
        width="100vw",
        height="100vh",
        bg=COLORS['background']
    )

# --- PÁGINAS ---

def page_generador():
    return rx.vstack(
        rx.heading("Generador de Escritos", color=COLORS['navy']),
        rx.divider(),
        
        rx.select(
            ["Prescripción de la Pena", "Extinción Art. 25 ter", "Amparo Constitucional", "Apelación por Quebrantamiento"],
            value=State.tipo_recurso,
            on_change=State.set_tipo_recurso,
            style=style_input
        ),
        
        rx.grid(
            rx.vstack(rx.text("Defensor/a", weight="bold"), rx.input(value=State.defensor_nombre, on_change=State.set_defensor_nombre, style=style_input)),
            rx.vstack(rx.text("Imputado/a", weight="bold"), rx.input(value=State.imputado, on_change=State.set_imputado, style=style_input)),
            columns="2",
            spacing="4",
            width="100%"
        ),
        
        rx.button("🧼 Limpiar Campos", on_click=State.clear_form, variant="outline", size="2"),
        
        # Renderizado condicional de formularios
        rx.cond(
            State.tipo_recurso == "Prescripción de la Pena",
            rx.vstack(
                rx.heading("Detalles Causa", size="4"),
                rx.input(placeholder="RIT", value=State.rit_input, on_change=State.set_rit_input, style=style_input),
                rx.input(placeholder="Pena", value=State.pena_input, on_change=State.set_pena_input, style=style_input),
                rx.button("➕ Agregar Causa", on_click=State.add_prescripcion_item, size="2"),
                rx.foreach(State.prescripcion_list, lambda x: rx.text(f"- {x['rit']}: {x['pena']}"))
            )
        ),
        
        rx.cond(
            State.tipo_recurso == "Apelación por Quebrantamiento",
            rx.vstack(
                rx.heading("Argumentación", size="4"),
                rx.text_area(placeholder="Hechos...", value=State.hechos_quebrantamiento, on_change=State.set_hechos_quebrantamiento, style=style_input),
                rx.text_area(placeholder="Argumentos...", value=State.argumentos_defensa, on_change=State.set_argumentos_defensa, style=style_input, height="150px"),
                rx.button("✨ Robustecer con IA", on_click=State.robustecer_argumentos, bg=COLORS['slate'], color="white", padding="0.5em"),
            )
        ),
        
        rx.button("🚀 GENERAR Y DESCARGAR DOCX", on_click=State.download_docx, style=style_button_primary, margin_top="2em"),
        width="100%",
        padding="2em",
        spacing="4",
        max_width="1000px"
    )

def page_analista():
    return rx.vstack(
        rx.heading("Analista Multimodal", color=COLORS['navy']),
        rx.text("Sube PDFs o Imágenes para análisis estratégico.", color=COLORS['slate']),
        
        rx.radio(
            ["Control de Detención", "Teoría del Caso", "Salidas Alternativas"],
            value=State.analisis_objetivo,
            on_change=State.set_analisis_objetivo,
            direction="row"
        ),
        
        rx.text_area(placeholder="Contexto adicional...", value=State.contexto_analisis, on_change=State.set_contexto_analisis, style=style_input),
        
        rx.upload(
            rx.vstack(
                rx.button("Seleccionar Archivos", bg=COLORS['white'], color=COLORS['navy'], border=f"1px solid {COLORS['navy']}"),
                rx.text("Arrastra archivos aquí", size="2")
            ),
            id="upload_analista",
            multiple=True,
            accept={"application/pdf": [".pdf"], "image/*": [".png", ".jpg"]},
            border=f"2px dashed {COLORS['beige']}",
            padding="2em",
        ),
        
        rx.button("⚡ ANALIZAR EVIDENCIA", on_click=State.handle_analisis_upload(rx.upload_files("upload_analista")), style=style_button_primary),
        
        rx.cond(
            State.is_analyzing,
            rx.center(rx.spinner(color=COLORS['navy'])),
            rx.cond(
                State.analisis_result != "",
                rx.card(
                    rx.markdown(State.analisis_result),
                    style=style_card,
                    width="100%"
                )
            )
        ),
        width="100%",
        padding="2em",
        spacing="4",
        max_width="1000px"
    )

def page_biblioteca():
    return rx.vstack(
        rx.heading("Biblioteca Jurídica RAG", color=COLORS['navy']),
        
        rx.flex(
            rx.select(["Todos", "Recurso de Nulidad", "Recurso de Amparo", "Sentencia"], value=State.filtro_tipo, on_change=State.set_filtro_tipo, style=style_input),
            rx.input(placeholder="Tribunal...", value=State.filtro_tribunal, on_change=State.set_filtro_tribunal, style=style_input),
            rx.input(placeholder="Tema Jurídico...", value=State.busqueda_query, on_change=State.set_busqueda_query, style=style_input, width="40%"),
            spacing="3",
            width="100%"
        ),
        
        rx.button("🔍 Investigar", on_click=State.buscar_jurisprudencia, style=style_button_primary),
        
        rx.cond(
            State.is_searching,
            rx.center(rx.spinner()),
            rx.vstack(
                rx.cond(
                    State.respuesta_juridica_ia != "",
                    rx.box(
                        rx.heading("⚖️ Respuesta Jurídica Inteligente", size="3"),
                        rx.markdown(State.respuesta_juridica_ia),
                        bg=COLORS['light_blue'],
                        padding="1.5em",
                        border_radius="10px",
                        border_left=f"5px solid {COLORS['navy']}",
                        width="100%"
                    )
                ),
                rx.divider(),
                rx.heading("Fuentes", size="3", color=COLORS['slate']),
                rx.foreach(
                    State.resultados_biblioteca,
                    lambda res: rx.card(
                        rx.flex(
                            rx.badge(res['meta_parsed']['tipo'], color_scheme="green"),
                            rx.badge(res['meta_parsed']['rol'], color_scheme="blue"),
                            rx.text(res['meta_parsed']['tribunal'], weight="bold"),
                            spacing="2"
                        ),
                        rx.text(res['contenido'], no_of_lines=3, size="2", margin_top="0.5em"),
                        style=style_card,
                        width="100%",
                        margin_bottom="1em"
                    )
                ),
                width="100%"
            )
        ),
        width="100%",
        padding="2em",
        spacing="4",
        max_width="1000px"
    )

def page_admin():
    return rx.vstack(
        rx.heading("Panel de Ingesta (Admin)", color=COLORS['navy']),
        rx.text("Sube documentos para alimentar el cerebro jurídico.", color=COLORS['slate']),
        
        rx.upload(
            rx.text("Arrastra PDFs aquí para indexar"),
            id="upload_ingesta",
            multiple=True,
            border=f"2px dashed {COLORS['beige']}",
            padding="3em"
        ),
        
        rx.button("💾 Procesar e Indexar", on_click=State.handle_ingesta_upload(rx.upload_files("upload_ingesta")), style=style_button_primary),
        
        rx.cond(
            State.is_ingesting,
            rx.spinner(),
            rx.text(State.admin_upload_status, color=COLORS['success'])
        ),
        width="100%",
        padding="2em",
        spacing="4"
    )

# --- LAYOUT PRINCIPAL ---

def index():
    return rx.cond(
        State.logged_in,
        rx.hstack(
            sidebar(),
            rx.box(
                rx.cond(State.current_page == "Generador", page_generador()),
                rx.cond(State.current_page == "Analista", page_analista()),
                rx.cond(State.current_page == "Biblioteca", page_biblioteca()),
                rx.cond(State.current_page == "Admin", page_admin()),
                width="100%",
                height="100vh",
                overflow_y="auto",
                bg=COLORS['background']
            ),
            spacing="0"
        ),
        login_page()
    )

# =============================================================================
# 5. INICIALIZACIÓN APP
# =============================================================================

app = rx.App(style={"font_family": "Segoe UI, sans-serif"})
app.add_page(index)